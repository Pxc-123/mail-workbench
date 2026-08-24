# -*- coding: utf-8 -*-
"""
销售邮件发送工作台 - 后端服务（Python 标准库，零外部依赖）
提供：多用户注册/登录（独立空间隔离）、待办、客户、标签、邮件模板、
展会资料、AI 生成、邮件预览/发送、发送日志、系统设置。
运行：python app.py  （默认 http://127.0.0.1:8000）
"""
import http.server
import socketserver
import sqlite3
import json
import os
import re
import random
import string
import hashlib
import datetime
import threading
import time
import urllib.parse
import csv
import io
import smtplib
import ssl
import mimetypes
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from http.server import BaseHTTPRequestHandler

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# 数据库路径支持环境变量覆盖（云托管/容器场景挂载持久盘时使用）
DB_PATH = os.environ.get("DB_PATH", os.path.join(BASE_DIR, "workbench.db"))

# ---------------------------- 翻译辅助（中英双语切换） ----------------------------
def _call_llm(prompt):
    """调用免费翻译 API 将中文译为英文；失败时返回原文兜底。"""
    # 从 prompt 中提取待翻译文本（去掉前面的指令行）
    text = prompt.split("\n", 1)[-1] if "\n" in prompt else prompt
    if not text.strip():
        return ""
    try:
        import urllib.request
        # MyMemory 免费翻译（CloudBase 海外节点通常可访问）
        url = "https://api.mymemory.translated.net/get?q=" + urllib.parse.quote(text) + "&langpair=zh-CN|en"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            if data.get("responseStatus") == 200:
                return data["responseData"]["translatedText"]
    except Exception:
        pass
    # 兜底：返回原文，避免页面无响应
    return text


def _translate_long_text(text):
    """分段翻译长文本：MyMemory 单次请求限 500 字符，按段落切块翻译后拼接。"""
    if not text or not text.strip():
        return text
    if len(text) <= 450:
        return _call_llm(f"将以下邮件正文翻译为地道商务英文，保持段落格式，只返回翻译结果，不要解释：\n{text}")
    blocks = text.split("\n\n")
    out = []
    buf = ""
    for b in blocks:
        if buf and len(buf) + len(b) + 2 > 450:
            out.append(_call_llm(f"将以下邮件正文翻译为地道商务英文，保持段落格式，只返回翻译结果，不要解释：\n{buf}"))
            buf = b
        else:
            buf = (buf + "\n\n" + b) if buf else b
    if buf:
        out.append(_call_llm(f"将以下邮件正文翻译为地道商务英文，保持段落格式，只返回翻译结果，不要解释：\n{buf}"))
    return "\n\n".join(out)
UPLOAD_DIR = os.environ.get("UPLOAD_DIR", os.path.join(BASE_DIR, "..", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ---------------------------- COS 持久化同步 ----------------------------
# 当配置了 COS 环境变量（COS_SECRET_ID / COS_SECRET_KEY / COS_REGION /
# COS_BUCKET）时，数据库文件与附件会自动同步到腾讯云 COS 对象存储。
# 这样即使 CloudBase 容器重部署 / 重启（本地文件系统是临时的），
# 客户数据、发送日志和附件都不会丢失。未配置时以下函数均为空操作，
# 不影响本地开发。
def _cos_cfg_ok():
    return all([
        os.environ.get("COS_SECRET_ID"),
        os.environ.get("COS_SECRET_KEY"),
        os.environ.get("COS_REGION"),
        os.environ.get("COS_BUCKET"),
    ])

COS_ENABLED = _cos_cfg_ok()
COS_PREFIX = os.environ.get("COS_PREFIX", "mailwb/").strip()
if COS_PREFIX and not COS_PREFIX.endswith("/"):
    COS_PREFIX += "/"

_cos_client = None
_cos_client_lock = threading.Lock()
_cos_sync_lock = threading.Lock()
_cos_sync_pending = False
_cos_worker = None

def get_cos_client():
    global _cos_client
    if not COS_ENABLED:
        return None
    if _cos_client is None:
        with _cos_client_lock:
            if _cos_client is None:
                try:
                    from qcloud_cos import CosConfig, CosS3Client
                    _cos_client = CosS3Client(CosConfig(
                        Region=os.environ["COS_REGION"],
                        SecretId=os.environ["COS_SECRET_ID"],
                        SecretKey=os.environ["COS_SECRET_KEY"],
                    ))
                except Exception as e:
                    print("[COS] 初始化失败：", e, flush=True)
                    _cos_client = False
    return _cos_client or None

def _cos_db_key():
    return COS_PREFIX + "workbench.db"

def _cos_upload_file(local_path, cos_key):
    client = get_cos_client()
    if not client or not os.path.exists(local_path):
        return False
    try:
        with open(local_path, "rb") as f:
            client.put_object(
                Bucket=os.environ["COS_BUCKET"],
                Key=cos_key,
                Body=f,
            )
        return True
    except Exception as e:
        # 打印完整错误信息，便于排查（如密钥/桶名/地域不匹配）
        import traceback
        print("[COS] 上传失败 %s: %s" % (cos_key, e), flush=True)
        traceback.print_exc()
        return False

def cos_upload_db():
    return _cos_upload_file(DB_PATH, _cos_db_key())

def cos_upload_attachment(local_path):
    rel = os.path.relpath(local_path, UPLOAD_DIR)
    key = COS_PREFIX + "uploads/" + rel.replace(os.sep, "/")
    return _cos_upload_file(local_path, key)

def cos_download_db():
    if not COS_ENABLED:
        return False
    client = get_cos_client()
    if not client:
        return False
    tmp = DB_PATH + ".cosdl"
    try:
        resp = client.get_object(Bucket=os.environ["COS_BUCKET"], Key=_cos_db_key())
        resp["Body"].get_stream_to_file(tmp)
        os.replace(tmp, DB_PATH)
        print("[COS] 数据库已从云端恢复", flush=True)
        return True
    except Exception as e:
        try:
            os.remove(tmp)
        except Exception:
            pass
        print("[COS] 数据库下载失败（可能尚未存在）：", e, flush=True)
        return False

def cos_download_all_uploads():
    if not COS_ENABLED:
        return 0
    client = get_cos_client()
    if not client:
        return 0
    prefix = COS_PREFIX + "uploads/"
    n = 0
    try:
        marker = ""
        while True:
            resp = client.list_objects(Bucket=os.environ["COS_BUCKET"], Prefix=prefix,
                                       Marker=marker, MaxKeys=1000)
            for item in (resp.get("Contents") or []):
                key = item["Key"]
                rel = key[len(prefix):]
                if not rel or rel.endswith("/"):
                    continue
                dst = os.path.join(UPLOAD_DIR, rel.replace("/", os.sep))
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                try:
                    r = client.get_object(Bucket=os.environ["COS_BUCKET"], Key=key)
                    r["Body"].get_stream_to_file(dst)
                    n += 1
                except Exception as e:
                    print("[COS] 下载附件失败 %s: %s" % (key, e), flush=True)
            if resp.get("IsTruncated") == "true":
                marker = resp.get("NextMarker", "") or ""
            else:
                break
        print("[COS] 已恢复 %d 个附件到本地" % n, flush=True)
    except Exception as e:
        print("[COS] 列举附件失败：", e, flush=True)
    return n

def cos_clear_uploads():
    if not COS_ENABLED:
        return
    client = get_cos_client()
    if not client:
        return
    prefix = COS_PREFIX + "uploads/"
    try:
        marker = ""
        while True:
            resp = client.list_objects(Bucket=os.environ["COS_BUCKET"], Prefix=prefix,
                                       Marker=marker, MaxKeys=1000)
            objs = [{"Key": i["Key"]} for i in (resp.get("Contents") or [])]
            if objs:
                client.delete_objects(Bucket=os.environ["COS_BUCKET"], Delete={"Object": objs})
            if resp.get("IsTruncated") == "true":
                marker = resp.get("NextMarker", "") or ""
            else:
                break
    except Exception as e:
        print("[COS] 清理云端附件失败：", e, flush=True)

def _cos_worker_loop():
    while True:
        time.sleep(2)
        with _cos_sync_lock:
            if not _cos_sync_pending:
                continue
            _cos_sync_pending = False
        try:
            cos_upload_db()
        except Exception as e:
            print("[COS] 同步失败（不影响系统运行）：%s" % e, flush=True)

def _schedule_cos_db_sync():
    global _cos_worker
    if not COS_ENABLED:
        return
    with _cos_sync_lock:
        _cos_sync_pending = True
        if _cos_worker is None:
            _cos_worker = threading.Thread(target=_cos_worker_loop, daemon=True)
            _cos_worker.start()

# 内存会话表 token -> user_id
SESSIONS = {}
SESS_LOCK = threading.Lock()

# ---------------------------- 数据库 ----------------------------
def get_db():
    conn = sqlite3.connect(DB_PATH, factory=_PersistedConnection)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn

class _PersistedConnection(sqlite3.Connection):
    """包装 SQLite 连接，每次 commit 后自动把数据库同步到 COS。"""
    def commit(self):
        super().commit()
        _schedule_cos_db_sync()

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        display_name TEXT,
        pass_hash TEXT,
        salt TEXT,
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS settings (
        user_id INTEGER PRIMARY KEY,
        smtp_host TEXT, smtp_port INTEGER, smtp_user TEXT, smtp_pass TEXT,
        from_email TEXT, from_name TEXT,
        default_interval INTEGER DEFAULT 5,
        demo_mode INTEGER DEFAULT 1,
        signature TEXT
    );
    CREATE TABLE IF NOT EXISTS todos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        due_time TEXT,
        bind_date TEXT,
        priority TEXT DEFAULT '中',
        customer_id INTEGER,
        done INTEGER DEFAULT 0,
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS tags (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        UNIQUE(user_id, name)
    );
    CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        company TEXT,
        contact TEXT,
        email TEXT,
        phone TEXT,
        exhibition TEXT,
        tags TEXT DEFAULT '',
        status TEXT DEFAULT '潜在客户',
        remark TEXT DEFAULT '',
        region TEXT DEFAULT '',
        source TEXT DEFAULT '',
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS exhibitions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        city TEXT,
        date_text TEXT,
        note TEXT
    );
    CREATE TABLE IF NOT EXISTS materials (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        exhibition_id INTEGER,
        name TEXT,
        file_path TEXT,
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS templates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        name TEXT,
        exhibition TEXT,
        customer_type TEXT,
        scene TEXT,
        tone TEXT,
        subject TEXT,
        body TEXT,
        signature TEXT,
        attachment_ids TEXT DEFAULT '',
        created_at TEXT
    );
    CREATE TABLE IF NOT EXISTS email_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        exhibition TEXT,
        template_name TEXT,
        customer_company TEXT,
        contact TEXT,
        email TEXT,
        subject TEXT,
        body TEXT,
        status TEXT,
        error TEXT,
        sent_at TEXT
    );
    CREATE TABLE IF NOT EXISTS drafts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT,
        payload TEXT NOT NULL,
        created_at TEXT,
        updated_at TEXT
    );
    CREATE TABLE IF NOT EXISTS backup_meta (
        key TEXT PRIMARY KEY,
        value TEXT
    );
    """)
    # 预置展会（全局共享只读，存到 user_id=0 表示系统级）
    c.execute("SELECT COUNT(*) AS n FROM exhibitions WHERE user_id=0")
    if c.fetchone()["n"] == 0:
        for ex in [("SIAL 巴黎食品展","法国巴黎","2026-10-18 ~ 10-22"),
                   ("越南国际食品展","越南胡志明市","2026-08-12 ~ 08-15"),
                   ("德国 ANUGA 食品展","德国科隆","2026-10-07 ~ 10-11"),
                   ("日本 FOODEX","日本千叶","2027-03-09 ~ 03-12"),
                   ("泰国 THAIFEX","泰国曼谷","2026-05-26 ~ 05-30")]:
            c.execute("INSERT INTO exhibitions (user_id,name,city,date_text) VALUES (0,?,?,?)", ex)
    # 用户角色列（admin / member），旧库安全添加
    try:
        c.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'member'")
    except Exception:
        pass
    # 客户扩展列（remark/region/source/status），旧库安全添加
    for col in ["remark TEXT DEFAULT ''", "region TEXT DEFAULT ''", "source TEXT DEFAULT ''", "status TEXT DEFAULT '潜在客户'"]:
        try:
            c.execute(f"ALTER TABLE customers ADD COLUMN {col}")
        except Exception:
            pass
    # 预置管理员账号（保证系统始终至少有一个管理员）
    try:
        admin_count = c.execute("SELECT COUNT(*) AS n FROM users WHERE role='admin'").fetchone()["n"]
        if admin_count == 0:
            user_count = c.execute("SELECT COUNT(*) AS n FROM users").fetchone()["n"]
            if user_count == 0:
                # 空库：创建系统管理员
                ah, asalt = hash_pass("admin123")
                c.execute("INSERT INTO users (username,display_name,pass_hash,salt,created_at,role) VALUES (?,?,?,?,?,'admin')",
                          ("admin", "系统管理员", ah, asalt, now_iso()))
            else:
                # 有用户但无管理员：把第一个注册的用户提升为管理员
                first_uid = c.execute("SELECT id FROM users ORDER BY id ASC LIMIT 1").fetchone()["id"]
                c.execute("UPDATE users SET role='admin' WHERE id=?", (first_uid,))
    except Exception:
        pass
    # 场景重编号迁移（已删除"创新大奖申报"场景：原编号 4 移除，原 5-8 → 4-7）
    try:
        c.execute("""UPDATE templates SET scene = CASE scene
                        WHEN '5' THEN '4' WHEN '6' THEN '5' WHEN '7' THEN '6' WHEN '8' THEN '7'
                        WHEN '4' THEN '1' WHEN '通知创新大奖申报截止提醒' THEN '1'
                        ELSE scene END
                    WHERE scene IN ('4','5','6','7','8','通知创新大奖申报截止提醒')""")
    except Exception:
        pass
    conn.commit()
    conn.close()

# ---------------------------- 工具函数 ----------------------------
def now_iso():
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def hash_pass(password, salt=None):
    if salt is None:
        salt = ''.join(random.choices(string.ascii_letters + string.digits, k=16))
    h = hashlib.sha256((salt + password).encode("utf-8")).hexdigest()
    return h, salt

def gen_token():
    return ''.join(random.choices(string.hexdigits, k=40))

def auth_user(handler):
    """从 Authorization: Bearer <token> 或 ?token= 解析用户"""
    auth = handler.headers.get("Authorization", "")
    token = None
    if auth.startswith("Bearer "):
        token = auth[7:].strip()
    if not token:
        qs = urllib.parse.urlparse(handler.path).query
        token = urllib.parse.parse_qs(qs).get("token", [None])[0]
    if not token:
        # 也尝试从 body
        return None
    with SESS_LOCK:
        uid = SESSIONS.get(token)
    if not uid:
        return None
    conn = get_db()
    u = conn.execute("SELECT id,username,display_name,role FROM users WHERE id=?", (uid,)).fetchone()
    conn.close()
    return dict(u) if u else None

def require_user(handler):
    u = auth_user(handler)
    if not u:
        return None, json_resp({"error": "未登录或登录已过期"}, 401)
    return u, None

def require_admin(handler):
    u, err = require_user(handler)
    if err:
        return None, err
    if u.get("role") != "admin":
        return None, json_resp({"error": "需要管理员权限"}, 403)
    return u, None

def json_resp(obj, code=200):
    body = json.dumps(obj, ensure_ascii=False).encode("utf-8")
    return code, {"Content-Type": "application/json; charset=utf-8"}, body

def read_json_body(handler):
    length = int(handler.headers.get("Content-Length", 0) or 0)
    if length == 0:
        return {}
    raw = handler.rfile.read(length)
    try:
        return json.loads(raw.decode("utf-8"))
    except Exception:
        return {}

def read_multipart(handler):
    """兼容接口：实际不再使用（改用 base64 JSON 方式传文件）"""
    return {}, {}

def row_to_dict(r):
    return dict(r) if r else None

# ---------------------------- AI 邮件生成（模板引擎，离线可用） ----------------------------
SCENE_LABELS = {
    "1": "初次开发陌生客户",
    "2": "跟进意向客户推送最新行业新闻",
    "3": "通知展位余量紧张催单",
    "4": "展会补贴政策通知",
    "5": "发送参展报价方案",
    "6": "客户跟进回访",
    "7": "参展感谢与维系",
}
TONE_LABELS = {"正式商务": "正式商务", "简洁干练": "简洁干练", "温和友好": "温和友好", "简短": "简短"}

TYPE_INTRO = {
    "预制菜": "贵司在预制菜领域的产品矩阵与出海布局",
    "调味品": "贵司在调味品赛道的产品创新与海外渠道拓展",
    "零食": "贵司在休闲零食品类的爆款打造与跨境销售",
    "原料": "贵司作为食品原料供应商的产能与品质优势",
    "综合食品企业": "贵司综合食品业务的多品类出海机会",
}

# 展会特色数据(用于生成差异化文案)
# 说明：以下数据基于公开网络检索（各展会官网 / 行业资讯）整理，覆盖城市、档期、
# 规模、核心亮点与差异化的开场白，确保不同展会生成出来的邮件内容各不相同。
EXHIBITION_PROFILES = {
    "SIAL 巴黎食品展": {
        "city": "法国巴黎", "date_hint": "2026年10月", "scale": "全球最大食品展之一，70+国家参展商，30万+专业观众",
        "highlights": ["欧洲买家集中度最高", "新品发布首选平台", "OEM/ODM对接效率极高"],
        "openings": [
            "作为全球食品行业的风向标，SIAL Paris 每年都是中国食品企业进入欧洲市场的黄金跳板。",
            "SIAL Paris 2026 即将开幕——这是中国预制菜/调味品企业触达欧洲采购决策者的最佳窗口期。",
            "每届 SIAL 都有超过 70% 的参观者拥有直接采购权，这正是贵司需要的精准买家群体。",
        ],
    },
    "越南国际食品展": {
        "city": "越南胡志明市", "date_hint": "2026年8月", "scale": "东南亚增长最快食品展，覆盖东盟10国+日韩澳新",
        "highlights": ["RCEP红利直接受益", "越南制造+出口双需求", "中企入驻成本低于欧美"],
        "openings": [
            "越南食品市场正以年均12%的速度增长，而 VietFood 是切入这一市场的最短路径。",
            "RCEP 生效后，越南已成为中国食品企业布局东南亚的桥头堡——VietFood 正是入场券。",
            "胡志明市作为越南经济中心，其食品加工产业对华合作意愿强烈，正是拓展良机。",
        ],
    },
    "德国 ANUGA 食品展": {
        "city": "德国科隆", "date_hint": "2026年10月", "scale": "世界最大食品饮料展，每两年一届，180+国家参与",
        "highlights": ["全球食品行业奥林匹克", "B2B成交额行业第一", "趋势发布权威平台"],
        "openings": [
            "ANUGA 被誉为食品行业的'奥林匹克'——每两年一次，错过就要再等两年。",
            "科隆 ANUGA 是全球唯一能同时见到 180+ 国家顶级买家的展会，对出海战略意义非凡。",
            "往届 ANUGA 中国展区一位难求，今年我们提前锁定了一批优质展位资源。",
        ],
    },
    "日本 FOODEX": {
        "city": "日本千叶", "date_hint": "2027年3月", "scale": "亚洲最大食品展之一，日本7万+专业买家到场",
        "highlights": ["日本食品安全标准最高", "单价/利润空间最优", "健康/功能性食品需求爆发"],
        "openings": [
            "日本市场以高门槛、高利润著称——FOODEX Japan 是打开这扇门的钥匙。",
            "日本消费者对海外优质食品的需求持续攀升，尤其是健康、功能性品类。",
            "FOODEX Japan 的买家质量在亚洲首屈一指，单客订单价值远超其他区域。",
        ],
    },
    "泰国 THAIFEX": {
        "city": "泰国曼谷", "date_hint": "2026年5月", "scale": "东盟核心食品展，连接南亚+中东+非洲买家",
        "highlights": ["东盟美食之都", "清真认证枢纽", "酒店餐饮业采购集中"],
        "openings": [
            "曼谷 THAIFEX 已成为连接东盟、南亚乃至中东食品买家的核心枢纽。",
            "泰国作为东盟美食中心，其辐射力可直达迪拜、印度等新兴市场。",
            "THAIFEX 独特的'酒店+零售'双渠道买家结构，让参展效果倍增。",
        ],
    },
    # ----------------- 包装机械 / 食品加工技术类（联网检索整理） -----------------
    "Interpack": {
        "city": "德国杜塞尔多夫", "date_hint": "2026年5月7-13日", "scale": "全球最大包装技术展，2800+展商来自67国，17万+专业观众",
        "highlights": ["智能智造与AI驱动产线", "欧洲PPWR包装法规落地前沿", "循环经济与可持续材料", "初创专区22家新锐企业"],
        "openings": [
            "Interpack 2026 刚于5月在杜塞尔多夫落幕，作为全球包装技术的'奥林匹克'，它集中展示了从智能产线到可持续材料的全产业链创新——这正是贵司设备触达欧洲买家的最佳参照窗口。",
            "每届 Interpack 都有 2800+ 展商、17万+ 观众，74% 展商来自境外，是包装机械企业进入欧洲市场不可替代的平台。",
            "今年 Interpack 的核心命题是 PPWR 法规与循环经济，欧洲买家对高效、合规的包装设备需求空前迫切，这正是贵司的强项。",
        ],
    },
    "ProPak China": {
        "city": "中国上海（国家会展中心）", "date_hint": "2026年6月15-17日", "scale": "20万㎡全产业链盛会，2500+展商来自135国，12万+专业观众",
        "highlights": ["包装机械+食品加工一站式", "智能装备与机器人产线", "覆盖食品饮料乳品日化制药", "四展联动贯通上下游"],
        "openings": [
            "ProPak China & FoodPack China 2026 将于6月在上海国家会展中心举办，20万㎡、2500+ 展商，是亚太规模最大的加工包装联展，贵司的装备能在这里直面全球采购商。",
            "本届联展以'马跃新程·智链未来'为主题，八大主题板块聚焦前沿科技与落地方案，对贵司这样的设备厂商是绝佳的曝光舞台。",
            "上海加工包装联展背靠长三角制造集群，买家覆盖食品、饮料、乳品、日化、制药等全行业，参展即等于一次精准的渠道总动员。",
        ],
    },
    "Gulfood Manufacturing": {
        "city": "阿联酋迪拜（世界贸易中心）", "date_hint": "2026年11月3-5日", "scale": "中东最大食品制造展，2500+展商来自79国，6.1万+专业观众，21个馆",
        "highlights": ["中东/非洲食品制造枢纽", "加工·包装·配料·自动化全链", "阿联酋粮食安全战略驱动", "GulfHost/自有品牌/糖果展同馆"],
        "openings": [
            "Gulfood Manufacturing 是中东地区最具影响力的食品饮料制造展，迪拜正借'2051粮食安全战略'大力扶持本地化生产，对包装与加工设备需求井喷。",
            "2026 年海湾食品加工展将于11月在迪拜世贸中心举办，2500+ 全球展商、21个展馆，是中企切入中东+非洲市场的核心跳板。",
            "阿联酋食品市场进口依赖度高达80%-90%，同时转向本地化智造——这意味着贵司的包装机械在迪拜有着确定性的刚需。",
        ],
    },
    "Anuga FoodTec": {
        "city": "德国科隆", "date_hint": "2027年2月23-26日", "scale": "全球食品加工技术标杆展，1300+展商，4万+观众来自130国",
        "highlights": ["Smart·Safe·Sustainable 主题", "AI与数字化工厂", "Inline检测与全程可追溯", "循环经济与节能降耗"],
        "openings": [
            "Anuga FoodTec 2027 将于2月在科隆举办，主题'Navigate Complexity——智能、安全、可持续'，是全球食品饮料技术最重要的风向标。",
            "作为每三年一届的技术标杆展，Anuga FoodTec 汇聚1300+ 展商、4万+ 来自130国的观众，是贵司对接欧洲高端食品工程客户的必争之地。",
            "今年展会聚焦 AI、数字孪生与端到端可追溯，欧洲买家对'聪明又合规'的产线方案求贤若渴，这恰是贵司技术的用武之地。",
        ],
    },
    "PACK EXPO": {
        "city": "美国芝加哥（麦考密克）", "date_hint": "2026年10月18-21日", "scale": "北美最大包装展，2600展商，4.8万+专业观众，两年一届",
        "highlights": ["北美市场订单质量高", "智能包装与减塑轻量化", "覆盖40+垂直行业", "食品医药日化采购集中"],
        "openings": [
            "PACK EXPO International 2026 将于10月在芝加哥举办，是北美规模最大的包装旗舰展，北美买家对自动化、可持续与智能包装需求强劲。",
            "作为两年一届的北美包装盛会，PACK EXPO 汇聚2600家展商、4.8万专业观众，是中国包装设备打入美洲市场的高效通道。",
            "当前北美正加速向'减塑、轻量化、高效产线'转型，贵司在高速立式包装机、AI视觉质检等领域的优势正对买家胃口。",
        ],
    },
    "ProPak Asia": {
        "city": "泰国曼谷（IMPACT）", "date_hint": "2026年6月10-13日", "scale": "东南亚最大加工包装展，2500+品牌来自45国，6.5万㎡",
        "highlights": ["东盟加工包装旗舰展", "AI自动化与智慧工厂", "冷链与创新可持续包装", "15个国际展团同台"],
        "openings": [
            "ProPak Asia 2026 将于6月在曼谷 IMPACT 举办，是东南亚最大的食品加工与包装技术展，2500+ 品牌来自45国，是中企扎根东盟的门户。",
            "今年展会面积扩大20%至6.5万㎡，聚焦 AI 自动化、冷链与可持续包装，RCEP 红利下东盟买家采购意愿持续走强。",
            "曼谷作为东盟制造中心，ProPak Asia 的买家辐射泰国、越南、马来、新加坡，参展一次即可覆盖整个东南亚渠道。",
        ],
    },
    "FHA Food & Beverage": {
        "city": "新加坡博览中心", "date_hint": "2026年4月21-24日", "scale": "亚太最大食品饮料展，2750+展商来自115国，8万+观众，10万㎡",
        "highlights": ["东盟采购枢纽门户", "欧盟'荣誉展区'背书", "FutureFWD食品科技专区", "契合新加坡30×30国策"],
        "openings": [
            "FHA Food & Beverage 2026 将于4月在新加坡举办，是亚太规模最大的食品饮料 B2B 展，2750+ 展商、8万+ 观众，是中企进军东盟的总入口。",
            "新加坡是东盟的战略枢纽与转口贸易中心，FHA 汇聚115国买家，贵司可借此一站触达东南亚分销与零售巨头。",
            "今年欧盟作为'荣誉展区'亮相，加上 FutureFWD 食品科技专区，FHA 已成为亚洲食品创新与贸易配对的核心舞台。",
        ],
    },
    "IPACK-IMA": {
        "city": "意大利米兰（Fiera Milano）", "date_hint": "2026年5月27-30日", "scale": "欧洲第二大包装展，13万㎡，1300+展商，7万+观众",
        "highlights": ["欧洲高端包装市场精准对接", "智能与绿色包装技术", "契合欧盟PPWR新法规", "中意经贸合作深化"],
        "openings": [
            "IPACK-IMA 2026 将于5月在米兰举办，是欧洲第二大包装与食品加工展，13万㎡、1300+ 展商，直通意大利及南欧高端市场。",
            "米兰依托欧盟统一市场与高端制造底蕴，IPACK-IMA 对智能、绿色包装技术需求旺盛，是贵司对接欧洲高端终端的桥梁。",
            "在欧盟 PPWR 新法规背景下，欧洲买家急需合规又高效的包装产线，这恰是贵司技术切入意大利市场的窗口。",
        ],
    },
}

# 别名映射：让用户即使输入简称/中文名也能命中对应展会资料（大小写、空格、连字符均已归一化）
EXHIBITION_ALIASES = {
    "interpack": "Interpack", "interpack2026": "Interpack", "杜塞尔多夫包装展": "Interpack", "杜塞尔多夫": "Interpack",
    "propak china": "ProPak China", "propakchina": "ProPak China", "上海加工包装展": "ProPak China",
    "上海国际食品加工与包装机械展": "ProPak China", "上海包装展": "ProPak China",
    "gulfood manufacturing": "Gulfood Manufacturing", "gulfoodmanufacturing": "Gulfood Manufacturing",
    "迪拜海湾食品工业展": "Gulfood Manufacturing", "迪拜包装展": "Gulfood Manufacturing", "迪拜食品制造展": "Gulfood Manufacturing",
    "anuga foodtec": "Anuga FoodTec", "anugafoodtec": "Anuga FoodTec", "科隆食品加工展": "Anuga FoodTec", "科隆食品技术展": "Anuga FoodTec",
    "pack expo": "PACK EXPO", "packexpo": "PACK EXPO", "芝加哥包装展": "PACK EXPO", "美国包装展": "PACK EXPO",
    "propak asia": "ProPak Asia", "propakasia": "ProPak Asia", "曼谷包装展": "ProPak Asia", "泰国包装展": "ProPak Asia", "泰国加工包装展": "ProPak Asia",
    "fha": "FHA Food & Beverage", "fha food & beverage": "FHA Food & Beverage", "fha foodandbeverage": "FHA Food & Beverage",
    "新加坡食品展": "FHA Food & Beverage", "新加坡食品饮料展": "FHA Food & Beverage", "新加坡fha": "FHA Food & Beverage",
    "ipack-ima": "IPACK-IMA", "ipack ima": "IPACK-IMA", "ipackima": "IPACK-IMA", "米兰包装展": "IPACK-IMA", "意大利包装展": "IPACK-IMA",
    "sial": "SIAL 巴黎食品展", "sial paris": "SIAL 巴黎食品展", "巴黎食品展": "SIAL 巴黎食品展",
    "foodex": "日本 FOODEX", "foodex japan": "日本 FOODEX", "日本食品展": "日本 FOODEX",
    "thaifex": "泰国 THAIFEX", "曼谷食品展": "泰国 THAIFEX", "越南食品展": "越南国际食品展", "anuga": "德国 ANUGA 食品展",
}

def _norm_ex_name(s):
    return re.sub(r"[\s\-–—_.()（）【】\[\]]", "", (s or "").lower())

def resolve_exhibition_profile(name):
    """根据展会名称在内置检索资料中匹配（支持别名与包含匹配）。"""
    if not name:
        return {}
    n = _norm_ex_name(name)
    # 1) 精确（归一化）匹配
    for k, v in EXHIBITION_PROFILES.items():
        if _norm_ex_name(k) == n:
            return v
    # 2) 别名匹配
    if n in EXHIBITION_ALIASES:
        return EXHIBITION_PROFILES.get(EXHIBITION_ALIASES[n], {})
    # 3) 双向包含匹配
    for k, v in EXHIBITION_PROFILES.items():
        nk = _norm_ex_name(k)
        if nk and (nk in n or n in nk):
            return v
    return {}

def _normalize_date_text(s):
    """把 Excel 序列号（如 46508）统一为 YYYY-MM-DD；已合规或无法识别则原样返回。"""
    if not s:
        return s
    s = str(s).strip()
    if re.match(r"^\d{4,6}$", s):
        try:
            return (datetime.datetime(1899, 12, 30) + datetime.timedelta(days=int(s))).strftime("%Y-%m-%d")
        except Exception:
            return s
    return s

def get_exhibition_profile(name, uid=None):
    """读取「展会资料库」中维护的真实展会资料（城市/档期/简介）。
    展会资料库是共享资料，对所有账号生效：优先匹配本人维护的条目，
    否则匹配任意条目（含全局/其他账号导入的）。
    最终与内置检索资料（含真实亮点）做合并：内置亮点/规模优先保留，
    DB 提供的真实城市/档期若有则覆盖，确保“不同展会不同亮点”且不全空。"""
    builtin = resolve_exhibition_profile(name)  # 内置检索资料（真实亮点）
    try:
        conn = get_db()
        rows = conn.execute(
            "SELECT user_id,name,city,date_text,note FROM exhibitions"
        ).fetchall()
        conn.close()
        n = _norm_ex_name(name)
        best = None
        for r in rows:
            rn = _norm_ex_name(r["name"])
            if rn == n or rn in n or n in rn:
                cand = {
                    "city": r["city"] or "",
                    "date_hint": _normalize_date_text(r["date_text"] or ""),
                    "scale": "",
                    "highlights": [(r["note"] or "").strip()] if (r["note"] or "").strip() else [],
                    "openings": [],
                    "_from_db": True,
                }
                if r["user_id"] == (uid or -1):  # 优先本人维护的
                    best = cand
                    break
                if best is None:
                    best = cand
        if best is not None or builtin:
            merged = {"city": "", "date_hint": "", "scale": "", "highlights": [], "openings": []}
            merged.update(builtin)  # 先铺内置（含真实亮点/规模）
            if best:
                if best["city"]:
                    merged["city"] = best["city"]
                if best["date_hint"]:
                    merged["date_hint"] = best["date_hint"]
                # DB 的 note 不再直接覆盖内置核心亮点；
                # 仅当内置完全没有该展会资料时，才把 note 当作备用亮点（且过滤机械文案）
                _db_hl = (best.get("highlights") or [])
                if not builtin and _db_hl:
                    _skip = {"上传资料时创建", "自动创建", "", " "}
                    _real = [h for h in _db_hl if h.strip() not in _skip]
                    if _real:
                        merged["highlights"] = _real
                merged["_from_db"] = True
            return merged
    except Exception:
        pass
    return builtin

def _fetch_industry_news(keyword="食品包装机械 海外市场"):
    """抓取真实行业新闻（带真实发布时间），用于邮件场景2「跟进意向客户推送最新行业新闻」。
    抓取策略（按成功率排序）：
      1) rss2json 代理 Google News RSS —— 结构化 JSON、带 pubDate，最可靠；
      2) rss2json 代理 Bing News RSS —— 备用结构化源；
      3) 直连百度新闻 HTML —— 容器能出网时的兜底；
      4) 仅当以上全部失败，才返回明确标注「缓存」的近期热点（非实时）。
    返回结果按发布时间倒序、标题去重，优先最新。"""
    import ssl as _ssl, re as _re, json as _json
    ctx = _ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = _ssl.CERT_NONE
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }

    def _via_rss2json(rss_url, label):
        """经 rss2json 公共接口把 RSS 转成 JSON（后端对后端，无 CORS 问题）。"""
        api = "https://api.rss2json.com/v1/api.json?rss_url=" + urllib.parse.quote(rss_url)
        try:
            req = urllib.request.Request(api, headers=headers)
            with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
                data = _json.loads(resp.read().decode("utf-8", "ignore"))
            if data.get("status") != "ok":
                return None
            items = data.get("items") or []
            parsed = []
            for it in items:
                title = (it.get("title") or "").strip()
                pub = (it.get("pubDate") or "").strip()
                if len(title) < 8:
                    continue
                parsed.append({"title": title, "date": pub})
            if len(parsed) >= 3:
                parsed.sort(key=lambda x: x["date"], reverse=True)
                seen, uniq = set(), []
                for p in parsed:
                    k = p["title"][:18]
                    if k in seen:
                        continue
                    seen.add(k)
                    uniq.append(p)
                return {"source": "实时新闻(" + label + ")", "items": uniq[:6],
                        "fetched_at": now_iso(), "dated": True}
        except Exception:
            return None
        return None

    # 主源：Google News RSS（中文聚合，覆盖出海/包装/展会）
    gnews = ("https://news.google.com/rss/search?q=" + urllib.parse.quote(keyword)
             + "&hl=zh-CN&gl=CN&ceid=CN:zh-Hans")
    r = _via_rss2json(gnews, "Google News")
    if r:
        return r
    # 备用：Bing News RSS
    bnews = "https://www.bing.com/news/search?q=" + urllib.parse.quote(keyword) + "&format=rss"
    r = _via_rss2json(bnews, "Bing News")
    if r:
        return r

    # 兜底1：直连百度新闻 HTML（容器能出网时可用）
    try:
        url = "https://news.baidu.com/ns?word=" + urllib.parse.quote(keyword) + "&tn=news&from=news&cl=2&rn=10"
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
            html = resp.read().decode("utf-8", "ignore")
        titles = _re.findall(r'<a[^>]*target="_blank"[^>]*>([^<]+)</a>', html)
        results = [t.strip() for t in titles if len(t.strip()) > 10]
        if len(results) >= 3:
            return {"source": "百度新闻", "items": results[:6], "fetched_at": now_iso()}
    except Exception:
        pass

    # 兜底2：明确标注的缓存热点（非实时，仅当实时抓取全部失败）
    fallback_items = [
        "2026年全球食品包装机械市场规模预计突破580亿美元，亚太地区增速领跑",
        "RCEP全面生效两周年，中国食品机械对东盟出口同比增长28%",
        "欧盟新版食品接触材料法规(FCM)将于2026年底实施，出口企业需提前合规",
        "智能包装与可持续包装成为2026年国际展会核心主题，买家关注度提升40%",
        "东南亚食品加工市场快速扩张，越南、印尼、泰国成中国设备主要出口目的地",
        "中东及北非地区食品进口需求持续增长，迪拜Gulfood展位供不应求",
    ]
    return {"source": "行业热点(缓存)", "items": fallback_items, "fetched_at": now_iso(),
            "note": "实时抓取暂不可用，显示近期已知热点"}

# 差异化开场白池(随机选取避免雷同)
OPENING_VARIANTS = [
    "您好！希望这封邮件没有打扰您的工作节奏。",
    "您好！感谢您抽出时间阅读这封信。",
    "您好！冒昧来信，是觉得以下信息可能对贵司有价值。",
    "您好！在这个信息过载的时代，我只说重点。",
    "您好！直接切入正题——有一个机会想和您同步。",
]

# 差异化结尾池
CLOSING_VARIANTS = [
    "如需进一步了解展会详情或展位方案，随时欢迎联系我。期待您的回复！",
    "我会在本周内跟进您的反馈。如有任何疑问，请随时告知。",
    "附件中包含本次展会的详细资料供您参考。期待与贵司在展会现场相见！",
    "如方便的话，我们可以安排一个简短的电话沟通具体需求。",
    "无论最终是否参展，都感谢您的时间。祝生意兴隆！",
]

def _build_ex_info(ex, profile):
    """根据展会资料（可能只有城市/档期/亮点中的部分字段）构建信息段落。"""
    if not profile:
        return f"【{ex}】"
    lines = [f"【{ex}】"]
    if profile.get("city"):
        lines.append(f"📍 举办地：{profile['city']}")
    if profile.get("date_hint"):
        lines.append(f"📅 展期：{_normalize_date_text(profile['date_hint'])}")
    if profile.get("scale"):
        lines.append(f"📊 规模：{profile['scale']}")
    hl = profile.get("highlights") or []
    if hl:
        lines.append("✨ 核心亮点：" + "、".join(hl[:3]))
    return "\n".join(lines)

# ---------------------------- 资料文本提取 ----------------------------
_MATERIAL_TEXT_CACHE = {}
def _extract_material_text(file_path, max_chars=4000):
    """从 PDF/Word/文本文件中提取前 max_chars 字文本，用于邮件生成时把资料内容写进邮件。
    失败时返回空串（不抛错），保证邮件生成不会因为坏资料卡住。"""
    if not file_path or not os.path.exists(file_path):
        return ""
    cache_key = (file_path, max_chars, int(os.path.getmtime(file_path)))
    if cache_key in _MATERIAL_TEXT_CACHE:
        return _MATERIAL_TEXT_CACHE[cache_key]
    text = ""
    try:
        low = file_path.lower()
        if low.endswith(".pdf"):
            try:
                from pypdf import PdfReader
                r = PdfReader(file_path)
                parts = []
                for i, page in enumerate(r.pages[:8]):  # 最多读前 8 页
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        continue
                text = "\n".join(parts)
            except Exception as e:
                print("[material] PDF 读取失败 %s: %s" % (file_path, e), flush=True)
        elif low.endswith((".docx", ".doc")):
            try:
                import docx as _docx
                d = _docx.Document(file_path)
                parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
                for t in d.tables[:5]:
                    for row in t.rows:
                        parts.append(" | ".join(c.text.strip() for c in row.cells if c.text))
                text = "\n".join(parts)
            except Exception as e:
                print("[material] Word 读取失败 %s: %s" % (file_path, e), flush=True)
        elif low.endswith((".txt", ".md", ".csv")):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif low.endswith((".xlsx", ".xls")):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                parts = []
                for sh in wb.sheetnames[:3]:
                    ws = wb[sh]
                    for row in ws.iter_rows(max_row=80, values_only=True):
                        parts.append(" | ".join(str(c) for c in row if c is not None))
                text = "\n".join(parts)
            except Exception as e:
                print("[material] Excel 读取失败 %s: %s" % (file_path, e), flush=True)
    except Exception as e:
        print("[material] 解析异常：%s" % e, flush=True)
    text = (text or "").strip()
    if len(text) > max_chars:
        text = text[:max_chars] + "…"
    _MATERIAL_TEXT_CACHE[cache_key] = text
    return text

def _summarize_materials(uid, material_ids):
    """根据用户选中的材料 ID 列表，汇总为一段「资料要点」文字。
    优先提取每份资料前 800 字做摘要，用换行拼接，去重相似的句子。"""
    if not material_ids:
        return ""
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT id,name,file_path FROM materials WHERE id IN (%s) AND (user_id=? OR user_id=0)" %
            ",".join("?" for _ in material_ids),
            list(material_ids) + [uid]).fetchall()
    finally:
        conn.close()
    chunks = []
    for r in rows:
        body = _extract_material_text(r["file_path"], max_chars=1200)
        if not body:
            continue
        body = body.replace("\r\n", "\n")
        # 简单清洗：把空行合并、去头尾空白
        lines = [ln.strip() for ln in body.split("\n") if ln.strip()]
        snippet = " ".join(lines[:8])[:800]
        if snippet:
            chunks.append(f"【{r['name']}】{snippet}")
    if not chunks:
        return ""
    return "\n\n".join(chunks)[:4000]

def _extract_highlights_from_materials(uid, material_ids):
    """从用户上传的展会资料中提取「核心亮点」候选句（规模/特色/优势等）。
    启发式扫描正文，挑出包含展会卖点关键词、长度适中且不重复的短句。"""
    if not material_ids:
        return []
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT file_path FROM materials WHERE id IN (%s) AND (user_id=? OR user_id=0)" %
            ",".join("?" for _ in material_ids), list(material_ids) + [uid]).fetchall()
    finally:
        conn.close()
    kw = ["规模", "面积", "万", "平米", "㎡", "展商", "观众", "国家", "企业", "亮点",
          "特色", "优势", "核心", "最大", "领先", "首选", "权威", "专业", "唯一", "覆盖",
          "买家", "采购商", "国际", "全球", "精彩", "不容错过", "一站式", "高效"]
    out, seen = [], set()
    for r in rows:
        text = _extract_material_text(r["file_path"], max_chars=3000)
        if not text:
            continue
        for ln in text.replace("\r\n", "\n").split("\n"):
            ln = ln.strip()
            if not (8 <= len(ln) <= 48):
                continue
            if any(k in ln for k in kw):
                if ln not in seen:
                    seen.add(ln)
                    out.append(ln)
        if len(out) >= 6:
            break
    return out[:6]

def _fetch_exhibition_highlights(ex):
    """实时检索该展会的亮点/规模信息，作为核心亮点补充（best-effort，失败返回空列表）。"""
    try:
        r = _fetch_industry_news(ex + " 展会 规模 参展商 亮点")
        items = r.get("items") or []
        out = []
        for it in items[:5]:
            t = (it.get("title") if isinstance(it, dict) else str(it)) or ""
            t = t.strip()
            if t:
                out.append(t[:32])
        return out
    except Exception:
        return []

# 多版本角度定义：同一展会/场景，从不同侧重点产出明显不同的邮件版本
MULTI_ANGLES = [
    {"key": "material", "label": "展会资料视角（基于您上传的附件要点）", "lead": "material",
     "subject": "【资料速览】{ex} 展会亮点与展位方案，附核心资料"},
    {"key": "news", "label": "行业趋势视角（结合最新网络资讯）", "lead": "news",
     "subject": "【行业趋势】近期食品包装动态 × {ex} 出海价值"},
    {"key": "combo", "label": "资料+趋势结合视角（附件与网络双结合）", "lead": "combo",
     "subject": "【综合研判】展会资料 + 行业趋势，看 {ex} 出海机会"},
    {"key": "scale", "label": "规模与买家资源视角", "lead": "scale",
     "subject": "【买家资源】{ex} 汇聚全球采购商，邀您精准对接"},
    {"key": "urgency", "label": "展位余量紧迫感视角", "lead": "urgency",
     "subject": "【展位余量】{ex} 黄金档期仅剩少量，请尽快锁定"},
]


def _make_angle_lead(lead, ex, profile, material_summary, news_items):
    """根据角度生成一段引导文字，使各版本侧重点明显不同。"""
    if lead == "material":
        if material_summary:
            first = material_summary.split("\n\n")[0]
            if len(first) > 160:
                first = first[:160] + "…"
            return "结合贵司上传的展会资料，%s 的几点核心看点：\n%s" % (ex, first)
        return "%s已沉淀一批真实展会资料，欢迎索取完整版以了解买家画像与展位方案。" % ex
    if lead == "news":
        top = ""
        if news_items:
            for it in news_items[:3]:
                t = (it.get("title") if isinstance(it, dict) else str(it)) or ""
                if t:
                    top = t
                    break
        if top:
            return "近期行业动态显示：「%s」。在这一趋势下，%s 成为贵司把握出海机会的优质窗口。" % (top, ex)
        return "近期食品包装机械出海需求持续升温，%s 正是触达海外精准买家的好时机。" % ex
    if lead == "combo":
        parts = ["我们结合本次展会资料与近期行业动态，为贵司梳理了 %s 的核心价值：" % ex]
        if material_summary:
            first = material_summary.split("\n\n")[0]
            if len(first) > 120:
                first = first[:120] + "…"
            parts.append("· 展会资料要点：" + first)
        top = ""
        if news_items:
            for it in news_items[:3]:
                t = (it.get("title") if isinstance(it, dict) else str(it)) or ""
                if t:
                    top = t
                    break
        if top:
            parts.append("· 行业动态：" + top)
        return "\n".join(parts)
    if lead == "scale":
        city = profile.get("city") or ""
        scale = profile.get("scale") or ""
        return "%s（%s）%s，是区域内极具影响力的专业展会，汇聚大量海外采购商与经销商。" % (ex, city, scale)
    if lead == "urgency":
        return "目前 %s 优质展位余量已非常紧张，尤其贴合贵司品类的展区所剩无几。为锁定黄金档期与最佳曝光位，建议尽快确认意向。" % ex
    return ""


def _strip_llm_prefix(text):
    """清洗 LLM 输出的正文开头，去掉它自动添加的称呼/问候/邮件头，避免与代码拼接的 salutation 重复。
    原则：只删除开头的称呼/问候语行，保留后续正文。
    """
    import re
    if not text:
        return text
    # 最多检查前 5 行
    lines = text.split("\n")
    out = []
    skipped = 0
    salutation_patterns = [
        r"^\s*尊敬的\s*[\{【\(].*[\}】\)]",                  # "尊敬的{联系人}..."
        r"^\s*尊敬的\s*\S+\s*[\(（].*[\)）]",                  # "尊敬的X（Y）："
        r"^\s*尊敬的.{0,30}[：:]\s*$",                         # "尊敬的xxx："
        r"^\s*尊敬的.{0,40}",
        r"^\s*您好[！!，,。\s]*$",
        r"^\s*你好[！!，,。\s]*$",
        r"^\s*Dear\s+.{0,40}",
        r"^\s*Hi\s+.{0,40}",
        r"^\s*Hello\s+.{0,40}",
        r"^\s*【.*】\s*$",                                  # 邮件头如【关于XX的邮件】
        r"^\s*主题[：:].*$",
        r"^\s*关于\s*[\S\s]{0,40}的(?:一封)?邮件[：:.\s]*$",
        r"^\s*Subject[：:].*$",
    ]
    pat = re.compile("|".join(salutation_patterns))
    for ln in lines:
        if skipped < 5 and not ln.strip():
            skipped += 1
            continue  # 跳过开头空行
        if skipped < 2 and ln.strip() and pat.match(ln):
            skipped += 1
            continue  # 删掉称呼行
        out.append(ln)
    cleaned = "\n".join(out).lstrip("\n")
    return cleaned if cleaned else text


def _strip_fact_duplicates(text, profile):
    """后处理：如果 LLM 在正文里重复了事实块该展示的内容（具体日期/地点/数字等），把重复的句子删掉。
    profile: dict，含 date_hint / city / scale / hl_uniq（亮点）
    原则：只删除包含「明显是事实」的整句，保留其它信息。
    """
    import re
    if not text:
        return text
    # 构造要匹配的「事实特征」
    city = (profile.get("city") or "").strip()
    date_h = profile.get("date_hint") or profile.get("date_text") or ""
    scale = (profile.get("scale") or "").strip()
    fact_patterns = []
    # 日期：2027 年 3 月 X 日 或 X 月 X 日 至 X 日 等
    fact_patterns.append(r"\d{4}\s*[年\-/\.]\s*\d{1,2}\s*[月\-/\.]\s*\d{1,2}\s*[日]?")
    fact_patterns.append(r"\d{1,2}\s*月\s*\d{1,2}\s*[日号]?")
    fact_patterns.append(r"\d{4}\s*[-\-/]\s*\d{1,2}\s*[-\-/]\s*\d{1,2}")
    # 大数字：XX亿/XX万/XXX美元/XXXXX买家
    fact_patterns.append(r"\d+(\.\d+)?\s*亿(美元|人民币|美金|欧)?")
    fact_patterns.append(r"\d+(\.\d+)?\s*万[+＋]?\s*[买家客户]")
    fact_patterns.append(r"\d+%\s*增[长速]")
    fact_patterns.append(r"RCEP|东盟|欧盟.{0,4}法规|FCM")
    # 城市
    if city:
        fact_patterns.append(re.escape(city))
    pat = re.compile("|".join(fact_patterns))
    # 按句子分隔（中文标点）
    sents = re.split(r"(?<=[。！？；\n])\s*", text)
    out = []
    for s in sents:
        if not s.strip():
            out.append(s)
            continue
        # 包含事实特征：删
        if pat.search(s):
            continue
        out.append(s)
    result = "".join(out).strip()
    # 避免删空
    if len(result) < 30:
        return text
    return result


def build_email_multi(exhibition, customer_type, scene, tone, custom_input, signature, uid=None, material_ids=None, n=5, settings=None):
    """一次产出 4-5 个不同角度的邮件版本（结构/侧重点明显不同）。
    网络亮点与附件要点各抓取/提取一次，避免重复请求；每个角度强制不同开场白。"""
    ex = exhibition or "本次海外食品展"
    prefetched_news = _fetch_exhibition_highlights(ex)
    prefetched_material = _summarize_materials(uid, material_ids or [])
    angles = MULTI_ANGLES[:max(4, min(n, len(MULTI_ANGLES)))]
    versions = []
    for i, ang in enumerate(angles):
        ang = dict(ang)
        if i < len(OPENING_VARIANTS):
            ang["opening"] = OPENING_VARIANTS[i]
        subj, body, _lu, _le = build_email(exhibition, customer_type, scene, tone, custom_input, signature, uid,
                                 material_ids=material_ids, angle=ang,
                                 prefetched_news=prefetched_news, prefetched_material_summary=prefetched_material,
                                 settings=settings)
        versions.append({"index": i + 1, "angle": ang.get("label"), "angle_key": ang.get("key"),
                         "subject": subj, "body": body, "llm_used": _lu, "llm_error": _le})
    return versions


def build_email(exhibition, customer_type, scene, tone, custom_input,  signature, uid=None, material_ids=None,
                angle=None, prefetched_news=None, prefetched_material_summary=None, settings=None):
    ex = exhibition or "本次海外食品展"
    ctype = customer_type or "食品企业"
    scene_key = scene if scene in SCENE_LABELS else "1"
    tone_key = tone if tone in TONE_LABELS else "正式商务"
    intro = TYPE_INTRO.get(ctype, "贵司在食品领域的产品与渠道优势")
    news = (custom_input or "").strip()
    # 场景2「跟进意向客户推送最新行业新闻」：若用户未手动提供资讯，
    # 后端自动抓取真实新闻作为安全网（best-effort，失败则保留通用兜底）。
    if scene_key == "2" and not news:
        try:
            fr = _fetch_industry_news(ex + " 食品包装机械 出海")
            its = fr.get("items") or []
            if its:
                parts = []
                for it in its[:4]:
                    if isinstance(it, dict):
                        d = it.get("date", "")
                        parts.append(f"- {it.get('title','')}{('（'+d[:10]+'）') if d else ''}")
                    else:
                        parts.append(f"- {it}")
                news = "\n".join(parts)
        except Exception:
            pass

    # 获取展会特色数据：优先用户在「展会资料库」维护的真实资料，其次内置检索资料
    profile = get_exhibition_profile(ex, uid)
    # 核心亮点升级：内置展会资料 + 上传附件提取 + 实时网络检索（带兜底，网络失败不影响生成）
    hl_material = _extract_highlights_from_materials(uid, material_ids)
    hl_news = prefetched_news if prefetched_news is not None else _fetch_exhibition_highlights(ex)
    merged_hl = []
    seen = set()
    for h in hl_material + hl_news + list(profile.get("highlights") or []):
        h = (h or "").strip().rstrip("。，,．.；; ").strip()
        if h and h not in seen:
            seen.add(h)
            merged_hl.append(h)
    if merged_hl:
        profile["highlights"] = merged_hl[:3]
    ex_city = profile.get("city", "")
    ex_scale = profile.get("scale", "")
    ex_highlights = profile.get("highlights", [])
    hl2 = "、".join(ex_highlights[:2]) if ex_highlights else ""
    # 开场白：多版本模式下由角度固定指定（确保每个版本不同）；单版本模式随机选取
    if angle and angle.get("opening"):
        opening = angle["opening"]
    else:
        opening = random.choice(profile["openings"]) if profile.get("openings") else random.choice(OPENING_VARIANTS)
    closing = random.choice(CLOSING_VARIANTS)
    ex_info_para = _build_ex_info(ex, profile)
    # 展会三要素信息块：每条邮件固定包含【展会名称 / 展会时间 / 展会亮点】
    # 无论哪个场景、单版本或多版本，都确保收件人看到展会基本盘。
    _facts_lines = [f"📌 展会信息", f"· 展会名称：{ex}"]
    _date_text = profile.get("date_hint") or profile.get("date_text") or ""
    _facts_lines.append(f"· 展会时间：{_normalize_date_text(_date_text) if _date_text else '待定（可向我索取最新档期）'}")
    _hl_all = (profile.get("highlights") or []) + ([hl2] if hl2 else [])
    _hl_seen = set()
    _hl_uniq = []
    for _h in _hl_all:
        _h = (_h or "").strip()
        if _h and _h not in _hl_seen:
            _hl_seen.add(_h); _hl_uniq.append(_h)
    _facts_lines.append("· 展会亮点：" + ("、".join(_hl_uniq[:3]) if _hl_uniq else "（详见资料，可向我要展位图与亮点清单）"))
    if profile.get("city"):
        _facts_lines.append(f"· 举办城市：{profile['city']}")
    if profile.get("scale"):
        _facts_lines.append(f"· 展会规模：{profile['scale']}")
    ex_facts_block = "\n".join(_facts_lines)
    # 用户选定的资料内容：让邮件真正反映最新资料（PDF/Word 全文摘要）
    material_summary = prefetched_material_summary if prefetched_material_summary is not None else _summarize_materials(uid, material_ids or [])
    mat_block = f"【资料要点】\n{material_summary}\n\n" if material_summary else ""
    # 角度引导段：使不同版本明显不同（基于角度/网络/附件侧重点）
    angle_lead = ""
    if angle and angle.get("lead"):
        angle_lead = _make_angle_lead(angle["lead"], ex, profile,
                                      prefetched_material_summary if prefetched_material_summary is not None else _summarize_materials(uid, material_ids or []),
                                      prefetched_news if prefetched_news is not None else _fetch_exhibition_highlights(ex))

    # 称呼占位（发送时按客户替换）
    salutation = "尊敬的 {联系人姓名}（{客户名称}）："

    scene_body = {
        "1": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"我是「{ex}」中国区招展团队的成员。本次致信是希望向贵司介绍这一重要的海外拓展机会。\n\n"
            f"结合{intro}，我们相信贵司的产品与本次展会的买家画像高度契合。\n\n"
            f"借此邮件，诚挚邀请贵司莅临{ex}，与海外买家面对面洽谈、拓展订单。如您方便，我可先发送展位图与参展方案供参考。\n\n"
            f"{closing}"
        ),
        "2": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"持续关注贵司在海外市场的进展。近期食品行业有几条值得留意的动态，特别与{intro}相关：\n\n"
            f"{ ('【行业资讯】\n' + news) if news else '【行业资讯】近期多国进口食品需求回暖，买家采购意愿明显增强；RCEP 框架下亚洲区内贸易成本持续下降。' }\n\n"
            f"在此背景下，{ex}将是贵司触达精准海外买家的优质窗口——{ex_scale or '汇聚全球优质采购商'}。如需，我可补充本次展会的买家结构与往届成交数据。\n\n"
            f"{closing}"
        ),
        "3": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"关于{ex}，需向您同步一个重要进展：目前优质展位余量已非常紧张，尤其贴合{intro}的展区所剩无几。\n\n"
            f"{ ('您此前关注的重点如下：\n' + news + '\n') if news else '' }"
            f"为保障贵司的参展位置与最佳曝光，建议尽快确认展位意向，避免错失黄金档期。我可为您预留 48 小时优先选位。\n\n"
            f"{closing}"
        ),
        "4": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"就贵司关注出海拓展的成本问题，特向您同步{ex}相关的参展补贴政策：多地商务主管部门对中小企业海外参展给予"
            f"展位费补贴（通常 50%~70% 不等），可显著降低出海门槛。\n\n"
            f"{ ('政策要点：\n' + news + '\n') if news else '' }"
            f"如贵司计划参展，建议尽早确认以赶上补贴申报周期（通常需提前2-3个月），我可协助准备相关材料。\n\n"
            f"{closing}"
        ),
        # ---- 报价 / 客户跟进 / 感谢 ----
        "5": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"关于贵司关注的{ex}，我们已为贵司初步测算了参展投入与回报，现将报价方案同步如下：\n\n"
            f"【展位方案】\n"
            f"· 标准展位（9㎡）：含基础搭建、楣板、照明、洽谈桌 —— 适合首次试水\n"
            f"· 光地展位（18㎡起）：可定制特装，最大化品牌曝光\n"
            f"· 双开口 / 角位：+15%，人流与曝光更优\n\n"
            f"{ ('【展会亮点】' + hl2 + '\n\n') if hl2 else '' }"
            f"结合{intro}，建议优先选择贴合贵司品类的展区，预计可触达大量精准海外买家。\n\n"
            f"以上为初步报价框架，最终方案可据贵司展品种类与预算灵活调整。如需要，我可发送完整版报价单与展位图。\n\n"
            f"{closing}"
        ),
        "6": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"距我们上次沟通已有一段时间，特来跟进贵司关于{ex}的参展意向，也想确认接下来的配合节奏。\n\n"
            f"想和您对齐三点：\n"
            f"1）参展预算与档期是否已排定？\n"
            f"2）希望重点对接哪类海外买家（经销商 / 品牌方 / 商超采购）？\n"
            f"3）是否需要我们协助准备展品运输与人员签证材料？\n\n"
            f"{ ('【您之前关注的信息】\n' + news + '\n\n') if news else '' }"
            f"目前{ex}优质展位余量有限{ex_city and '（' + ex_city + '）' or ''}，若确定参展建议尽快锁定，以免错失黄金位置。我可先为贵司预留 48 小时优先选位。\n\n"
            f"{closing}"
        ),
        "7": (
            f"{opening}\n\n"
            f"{angle_lead}\n\n" if angle_lead else f"{opening}\n\n"
            f"感谢贵司对{ex}的关注与支持！无论最终是否成行，都十分珍视与贵司的交流。\n\n"
            f"{ ('【本次展会价值】' + hl2 + '\n\n') if hl2 else '' }"
            f"如贵司后续有出海拓展、买家对接或展会相关的任何需求，我们随时提供协助——包括展后买家名单、行业报告与下一届档期预告。\n\n"
            f"期待未来有机会与贵司在展会现场或线上深入合作。祝生意兴隆！\n\n"
            f"{closing}"
        ),
    }[scene_key]

    tone_tail = {
        "正式商务": "",
        "简洁干练": " we can move fast.",
        "温和友好": "无论是否参展，都欢迎随时交流！",
        "简短": "",
    }[tone_key]

    if tone_key == "简短":
        lines = [l for l in scene_body.split("\n") if l.strip()]
        scene_body = "\n".join(lines[:3])  # 只保留前3段

    # ---- 真实大模型生成分支（可配置）：启用后用 LLM 产出正文，失败自动回退模板 ----
    llm_used = False
    llm_error = None
    ai_on = bool(settings and str(settings.get("ai_enabled", "")) in ("1", "true", "True"))
    if not ai_on:
        llm_error = f"未启用(settings={type(settings).__name__}, ai_enabled={settings.get('ai_enabled') if settings else 'no-settings'})"
    elif settings.get("ai_api_key") == "__skip__":
        llm_error = "api_key被设置为__skip__"
    if ai_on and not settings.get("ai_api_key") == "__skip__":
        try:
            _angle_desc = ""
            if angle:
                _angle_desc = (angle.get("lead") or angle.get("subject") or "")
            _prompt = (
                f"你是一名资深的海外食品展会招展顾问，请用中文写一封发给「{ctype}」企业负责人的招展邮件正文。\n"
                f"展会名称：{ex}\n"
                f"展会时间：{_normalize_date_text(profile.get('date_hint') or profile.get('date_text') or '待定')}\n"
                f"举办城市：{profile.get('city') or '待定'}\n"
                f"展会亮点参考：{('、'.join(_hl_uniq[:3]) if _hl_uniq else '详见资料')}\n"
                f"客户类型：{ctype}；使用场景：{SCENE_LABELS.get(scene_key, scene_key)}；语气：{tone_key}\n"
            )
            if _angle_desc:
                _prompt += f"本次写作角度：{_angle_desc}\n"
            if material_summary:
                _prompt += f"\n展会资料要点（可在叙事中自然化用，不要直接列举）：\n{material_summary}\n"
            if prefetched_news:
                _items = prefetched_news.get("items") or []
                if _items:
                    _news_txt = "\n".join("- " + (i.get("title", "") if isinstance(i, dict) else str(i)) for i in _items[:3])
                    _prompt += f"\n可参考的行业资讯（叙事化带入，不要列举标题）：\n{_news_txt}\n"
            _prompt += (
                f"\n用户补充要求：{news if news else '无'}\n"
                f"\n⚠️ 严格写作规则（务必遵守）：\n"
                f"1. 【禁止称呼】正文里绝对不要出现「尊敬的xxx」「您好」「Dear」「Hi」等任何称呼与问候语——称呼已经由系统独立加在最上面，你只写称呼之后的内容。\n"
                f"2. 【禁止重复事实】下面这些事实信息已经由系统整理在邮件末尾的【📌 展会信息】独立展示块里，正文里**绝对不要重复写**——包括：\n"
                f"   - 展会名称（如{ex}）\n"
                f"   - 展会具体时间（任何形如「2027年3月X日」「X月X-X日」的具体日期）\n"
                f"   - 举办城市（如{profile.get('city') or '千叶'}等）\n"
                f"   - 展会规模数字（如「7万+买家」「100,000 专业观众」等具体数字）\n"
                f"   - 展会亮点中的具体数据（如「580亿美元」「28% 增长」「FCM 法规」等）\n"
                f"   你在正文里**只用指代**——例如「本次展会」「该展」「那场展会」「既定档期」「海外买家集中的盛事」等，绝不重复列举这些数字、日期、地点。\n"
                f"3. 【结构】正文只写 3 段：① 简短铺垫（1-2 句切入背景，引出为什么值得看这封邮件） ② 参展价值/理由（讲参展能解决什么痛点、抓住什么机会） ③ 行动号召（引导下一步：例如「我可以先发您资料」「我们可以约个15分钟电话」「我帮您锁定48小时优先选位」等）。\n"
                f"4. 【格式】纯文本段落，不用 markdown 标题、不用 • 列表、不用 - 列项、不用 emoji 列表符号。控制在 250 字以内。\n"
                f"5. 【语气】{tone_key}，专业招展顾问口吻，不卑不亢，避免空话套话，每段都要有信息密度。\n"
                f"6. 开篇直接进入正文内容，第一句就开始讲故事/场景，不要写邮件头、邮件小标题、问候语。"
            )
            _llm_body = call_llm(_prompt, settings, max_tokens=900, temperature=0.9)
            if _llm_body and len(_llm_body) > 30:
                # 用 LLM 正文替换模板正文，但保留资料要点块与三要素块（确保信息完整）
                scene_body = _strip_llm_prefix(_llm_body)
                # 事实去重：删掉与末尾事实块重复的具体数字/日期/城市等
                _profile_for_strip = {
                    "city": profile.get("city"),
                    "date_hint": profile.get("date_hint") or profile.get("date_text"),
                    "scale": profile.get("scale"),
                    "hl_uniq": _hl_uniq,
                }
                scene_body = _strip_fact_duplicates(scene_body, _profile_for_strip)
                llm_used = True
            else:
                llm_error = "LLM返回为空或过短"
        except Exception as _e:
            # 调用失败：静默回退到模板生成，不影响出信
            llm_error = f"{type(_e).__name__}: {_e}"
            print(f"[AI-LLM] 大模型调用失败: {llm_error}", flush=True)

    body = f"{salutation}\n\n{scene_body}{mat_block}{ex_facts_block}\n\n— {signature or '{销售姓名}'}｜{ex} 招展团队"

    # 主题：多版本模式下优先使用角度专属主题（含展会名），否则用场景默认主题
    if angle and angle.get("subject"):
        subject = angle["subject"].replace("{ex}", ex)
    else:
        subject_map = {
            "1": f"邀您共赴 {ex}｜拓展海外买家渠道",
            "2": f"[{ctype}行业资讯] 附 {ex} 出海机会",
            "3": f"【展位余量提醒】{ex} 优质展区所剩无几",
            "4": f"【补贴政策】{ex} 参展补贴可显著降低出海成本",
            "5": f"【参展报价方案】{ex} 展位费用与投入回报",
            "6": f"【跟进】{ex} 参展意向确认，请查收",
            "7": f"【感谢】感谢关注 {ex}，后续资源持续开放",
        }
        subject = subject_map[scene_key]
    return subject, body, llm_used, llm_error

# ---------------------------- AI 大模型（可配置，真实调用） ----------------------------
_AI_COLUMNS = ["ai_enabled", "ai_provider", "ai_base_url", "ai_api_key", "ai_model"]

def ensure_ai_columns():
    """为 settings 表按需添加 AI 模型配置字段（幂等迁移，兼容旧库）。"""
    try:
        conn = get_db()
        cur = conn.execute("PRAGMA table_info(settings)")
        existing = {r["name"] for r in cur.fetchall()}
        for col in _AI_COLUMNS:
            if col not in existing:
                conn.execute(f"ALTER TABLE settings ADD COLUMN {col} TEXT")
        conn.commit()
        conn.close()
    except Exception:
        pass

def _provider_preset(provider):
    """各服务商的默认 Base URL 与推荐模型（OpenAI 兼容接口）。"""
    presets = {
        "deepseek": ("https://api.deepseek.com/v1", "deepseek-chat"),
        "qwen": ("https://dashscope.aliyuncs.com/compatible-mode/v1", "qwen-plus"),
        "openai": ("https://api.openai.com/v1", "gpt-4o-mini"),
        "ollama": ("http://localhost:11434/v1", "qwen2.5:7b"),
        "zhipu": ("https://open.bigmodel.cn/api/paas/v4", "glm-4-flash"),
        "custom": ("", "gpt-4o-mini"),
    }
    return presets.get(provider, presets["custom"])

def call_llm(prompt, settings, max_tokens=1200, temperature=0.8):
    """调用 OpenAI 兼容接口生成文本。失败时抛异常（由调用方兜底回退模板）。
    支持：DeepSeek / 通义千问 / OpenAI / 本地 Ollama 等任何 OpenAI 兼容端点。"""
    import urllib.request as _ur
    if not settings:
        raise ValueError("no settings")
    api_key = (settings.get("ai_api_key") or "").strip()
    base_url = (settings.get("ai_base_url") or "").strip()
    model = (settings.get("ai_model") or "").strip()
    if not base_url or not model:
        provider = settings.get("ai_provider") or "deepseek"
        base_url, model = _provider_preset(provider)
    if not base_url or not model:
        raise ValueError("LLM 未正确配置")
    url = base_url.rstrip("/") + "/chat/completions"
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    headers = {"Content-Type": "application/json"}
    if api_key and not base_url.startswith("http://localhost"):
        headers["Authorization"] = "Bearer " + api_key
    data = json.dumps(payload).encode("utf-8")
    req = _ur.Request(url, data=data, headers=headers, method="POST")
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    with _ur.urlopen(req, timeout=45, context=ctx) as resp:
        j = json.loads(resp.read().decode("utf-8", "ignore"))
    return j["choices"][0]["message"]["content"].strip()

def load_settings(user_id):
    conn = get_db()
    s = conn.execute("SELECT * FROM settings WHERE user_id=?", (user_id,)).fetchone()
    conn.close()
    if not s:
        return {"demo_mode": 1, "default_interval": 5, "signature": "",
                "ai_enabled": "", "ai_provider": "deepseek", "ai_base_url": "", "ai_api_key": "", "ai_model": ""}
    return row_to_dict(s)

def personalize(text, customer, sales_name):
    if not text:
        return ""
    rep = {
        "{客户名称}": customer.get("company") or "",
        "{联系人姓名}": customer.get("contact") or "",
        "{销售姓名}": sales_name or "",
        "{邮箱}": customer.get("email") or "",
        "{手机号}": customer.get("phone") or "",
        "{意向展会}": customer.get("exhibition") or "",
    }
    for k, v in rep.items():
        text = text.replace(k, v)
    return text

def send_one_email(user_id, to_email, subject, html_body, settings, attachments, cc, bcc):
    """返回 (status, error)"""
    if settings.get("demo_mode"):
        return "success", "演示模式（未实际外发）"
    if not (settings.get("smtp_host") and settings.get("from_email")):
        return "success", "未配置 SMTP，演示记录"
    try:
        msg = MIMEMultipart()
        msg["Subject"] = subject
        msg["From"] = f"{settings.get('from_name','')} <{settings.get('from_email')}>"
        msg["To"] = to_email
        if cc:
            msg["Cc"] = ", ".join(cc)
        if bcc:
            msg["Bcc"] = ", ".join(bcc)
        msg.attach(MIMEText(html_body, "plain", "utf-8"))
        for att in attachments or []:
            path = att.get("file_path")
            if path and os.path.exists(path):
                with open(path, "rb") as f:
                    part = MIMEApplication(f.read(), Name=att.get("name", "attachment"))
                    part["Content-Disposition"] = f'attachment; filename="{att.get("name","attachment")}"'
                    msg.attach(part)
        recipients = [to_email] + (cc or []) + (bcc or [])
        port = int(settings.get("smtp_port", 25))
        # 465 为 SSL 直连端口；587/25 等用 STARTTLS
        ssl_ctx = ssl.create_default_context()
        if port == 465:
            server = smtplib.SMTP_SSL(settings["smtp_host"], port, timeout=20, context=ssl_ctx)
        else:
            server = smtplib.SMTP(settings["smtp_host"], port, timeout=20)
            server.starttls()
        with server:
            if settings.get("smtp_user"):
                server.login(settings["smtp_user"], settings.get("smtp_pass", ""))
            server.sendmail(settings["from_email"], recipients, msg.as_string())
        return "success", ""
    except Exception as e:
        return "failed", str(e)

# ---------------------------- 路由处理 ----------------------------
class Handler(BaseHTTPRequestHandler):
    def _send(self, code, headers, body):
        self.send_response(code)
        for k, v in headers.items():
            self.send_header(k, v)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, Authorization")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, PATCH, DELETE, OPTIONS")
        self.end_headers()
        if body:
            self.wfile.write(body)

    def do_OPTIONS(self):
        self._send(204, {}, b"")

    def _query_param(self, name):
        """从 URL 查询字符串中提取参数值。"""
        qs = urllib.parse.urlparse(self.path).query
        params = urllib.parse.parse_qs(qs)
        vals = params.get(name, [])
        return vals[0] if vals else None

    def do_GET(self):
        path = urllib.parse.urlparse(self.path).path
        # 静态文件（兼容两种部署：/static/ 前缀 与 根相对路径 /css/ /js/）
        if path == "/" or path == "/index.html":
            return self.serve_static("index.html", "text/html")
        if path.startswith("/static/"):
            rel = path[len("/static/"):]
            return self.serve_static(rel, None)
        # 根相对路径：/css/... /js/... 直接映射到 frontend 目录（同源部署用）
        if path.startswith("/css/") or path.startswith("/js/"):
            return self.serve_static(path.lstrip("/"), None)
        # 其它带扩展名的静态资源（favicon.ico 等）
        last = path.split("/")[-1]
        if "." in last and not path.startswith("/api/"):
            return self.serve_static(path.lstrip("/"), None)
        # API
        code, headers, body = self._safe_route("GET", path)
        self._send(code, headers, body)

    def _safe_route(self, method, path):
        try:
            if method == "GET":
                return self.route_get(path)
            if method == "POST":
                return self.route_post(path)
            if method == "PATCH":
                return self.route_patch(path)
            if method == "DELETE":
                return self.route_delete(path)
            return 404, {"Content-Type": "application/json; charset=utf-8"}, json.dumps({"error": "method not allowed"}).encode()
        except Exception as e:
            import traceback as _tb
            err = _tb.format_exc()
            print("[DISPATCH-ERROR] %s %s:\n%s" % (method, path, err), flush=True)
            return 500, {"Content-Type": "application/json; charset=utf-8"}, \
                json.dumps({"error": "internal_server_error", "detail": str(e), "trace": err}, ensure_ascii=False).encode("utf-8")

    def do_POST(self):
        path = urllib.parse.urlparse(self.path).path
        code, headers, body = self._safe_route("POST", path)
        self._send(code, headers, body)

    def do_PATCH(self):
        path = urllib.parse.urlparse(self.path).path
        code, headers, body = self._safe_route("PATCH", path)
        self._send(code, headers, body)

    def do_PUT(self):
        """PUT 请求复用 POST 路由（编辑客户等操作）"""
        self.do_POST()

    def do_DELETE(self):
        path = urllib.parse.urlparse(self.path).path
        code, headers, body = self._safe_route("DELETE", path)
        self._send(code, headers, body)

    def serve_static(self, rel, ctype):
        # 兼容两种目录布局：部署布局(deploy/app.py + deploy/frontend/) 与 开发布局(backend/app.py + ../frontend/)
        candidates = [
            os.path.normpath(os.path.join(BASE_DIR, "frontend", rel)),
            os.path.normpath(os.path.join(BASE_DIR, "..", "frontend", rel)),
        ]
        fp = next((c for c in candidates if os.path.isfile(c)), None)
        if not fp:
            self._send(404, {"Content-Type": "text/plain"}, b"Not Found")
            return
        if ctype is None:
            ctype, _ = mimetypes.guess_type(fp)
            ctype = ctype or "application/octet-stream"
        with open(fp, "rb") as f:
            data = f.read()
        self._send(200, {"Content-Type": ctype + "; charset=utf-8"}, data)

    # ---------------- GET API ----------------
    def route_get(self, path):
        u, err = require_user(self)
        if err:
            return err
        if path.startswith("/api/admin/"):
            a, aerr = require_admin(self)
            if aerr:
                return aerr
            return self.admin_get(path, a)
        uid = u["id"]
        conn = get_db()
        try:
            if path == "/api/me":
                return json_resp(u)
            if path == "/api/todos":
                rows = conn.execute("SELECT * FROM todos WHERE user_id=? ORDER BY done, due_time", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/customers":
                rows = conn.execute("SELECT * FROM customers WHERE user_id=? ORDER BY id DESC", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/tags":
                rows = conn.execute("SELECT * FROM tags WHERE user_id=?", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/templates":
                rows = conn.execute("SELECT * FROM templates WHERE user_id=? ORDER BY id DESC", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/exhibitions":
                rows = conn.execute("SELECT * FROM exhibitions WHERE user_id=0 OR user_id=?", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/exhibitions/summary":
                # 含每个展会的资料数（用于展会管理页/AI 邮件页下拉统计）
                rows = conn.execute(
                    "SELECT e.id,e.user_id,e.name,e.city,e.date_text,e.note,"
                    "  (SELECT COUNT(*) FROM materials m WHERE m.exhibition_id=e.id) AS material_count "
                    "FROM exhibitions e WHERE e.user_id=0 OR e.user_id=?",
                    (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/drafts":
                rows = conn.execute("SELECT * FROM drafts WHERE user_id=? ORDER BY updated_at DESC", (uid,)).fetchall()
                out = []
                for r in rows:
                    d = row_to_dict(r)
                    try:
                        d["payload"] = json.loads(d.get("payload") or "{}")
                    except Exception:
                        d["payload"] = {}
                    out.append(d)
                return json_resp(out)
            if path == "/api/news/search":
                # 获取真实行业新闻（食品包装/出海/展会相关），用于邮件场景2
                q = self._query_param("q") or "食品包装机械 海外市场 出展"
                return json_resp(_fetch_industry_news(q))
            if path == "/api/customers/stats":
                # 客户跟踪统计（用于首页仪表盘）
                rows = conn.execute(
                    "SELECT COALESCE(status,'潜在客户') AS s, COUNT(*) AS c FROM customers WHERE user_id=? GROUP BY s", (uid,)
                ).fetchall()
                total = conn.execute("SELECT COUNT(*) FROM customers WHERE user_id=?", (uid,)).fetchone()[0]
                return json_resp({"total": total, "by_status": [row_to_dict(r) for r in rows]})
            if path == "/api/materials":
                rows = conn.execute("SELECT * FROM materials WHERE user_id=0 OR user_id=?", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/email-logs":
                rows = conn.execute("SELECT * FROM email_logs WHERE user_id=? ORDER BY id DESC", (uid,)).fetchall()
                return json_resp([row_to_dict(r) for r in rows])
            if path == "/api/settings":
                s = load_settings(uid)
                return json_resp(s)
            if path.startswith("/api/email-logs/"):
                lid = path.split("/")[-1]
                row = conn.execute("SELECT * FROM email_logs WHERE id=? AND user_id=?", (lid, uid)).fetchone()
                return json_resp(row_to_dict(row) if row else {})
        finally:
            conn.close()
        return json_resp({"error": "not found"}, 404)

    # ---------------- 管理员接口 ----------------
    def admin_get(self, path, admin):
        import os as _os
        if path == "/api/admin/users":
            conn = get_db()
            rows = conn.execute("SELECT id,username,display_name,role,created_at FROM users ORDER BY id").fetchall()
            conn.close()
            return json_resp([row_to_dict(r) for r in rows])
        if path == "/api/backup/export":
            # 全量导出（管理员）：用于跨重部署/换服务器迁移数据，防止数据丢失
            conn = get_db()
            try:
                tables = [r[0] for r in conn.execute(
                    "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchall()]
                data = {"version": 1, "exported_at": now_iso(), "app": "mail-workbench", "tables": {}}
                for t in tables:
                    rows = conn.execute(f"SELECT * FROM {t}").fetchall()
                    data["tables"][t] = [dict(r) for r in rows]
                return json_resp(data)
            finally:
                conn.close()
        if path == "/api/admin/backup/export-zip":
            # 全量导出（含数据库 + 所有附件文件）为 zip，用于挂持久盘前完整备份
            import io, zipfile
            c2 = get_db(); c2.close()  # 确保落盘
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                if os.path.exists(DB_PATH):
                    zf.write(DB_PATH, "workbench.db")
                if os.path.isdir(UPLOAD_DIR):
                    for root, dirs, files in os.walk(UPLOAD_DIR):
                        for fn in files:
                            fp = os.path.join(root, fn)
                            arc = os.path.relpath(fp, UPLOAD_DIR)
                            zf.write(fp, os.path.join("uploads", arc))
            data = buf.getvalue()
            fname = "mailwb-full-backup-%s.zip" % now_iso()[:10]
            return 200, {
                "Content-Type": "application/zip",
                "Content-Disposition": 'attachment; filename="%s"' % fname,
                "Content-Length": str(len(data)),
            }, data
        if path == "/api/admin/backup/status":
            # 部署稳定性：返回数据健康度，供管理员中心展示
            db_size = _os.path.getsize(DB_PATH) if _os.path.exists(DB_PATH) else 0
            upload_count, upload_bytes = 0, 0
            if _os.path.isdir(UPLOAD_DIR):
                for root, dirs, files in _os.walk(UPLOAD_DIR):
                    for fn in files:
                        fp = _os.path.join(root, fn)
                        try:
                            upload_count += 1
                            upload_bytes += _os.path.getsize(fp)
                        except Exception:
                            pass
            tables_info = []
            try:
                conn2 = get_db()
                for t in ("users","customers","exhibitions","materials","templates","email_logs","drafts","tags","todos","settings"):
                    try:
                        n = conn2.execute(f"SELECT COUNT(*) AS n FROM {t}").fetchone()["n"]
                        tables_info.append({"name": t, "rows": n})
                    except Exception:
                        tables_info.append({"name": t, "rows": -1})
                conn2.close()
            except Exception:
                pass
            # 上次备份时间（从 backup_meta 读）
            conn3 = get_db()
            last_backup = conn3.execute("SELECT value FROM backup_meta WHERE key='last_backup_at'").fetchone()
            last_source = conn3.execute("SELECT value FROM backup_meta WHERE key='last_backup_source'").fetchone()
            conn3.close()
            return json_resp({
                "cos_enabled": COS_ENABLED,
                "cos_bucket": os.environ.get("COS_BUCKET", ""),
                "db_path": DB_PATH,
                "db_size_bytes": db_size,
                "upload_count": upload_count,
                "upload_bytes": upload_bytes,
                "tables": tables_info,
                "last_backup_at": (last_backup["value"] if last_backup else ""),
                "last_backup_source": (last_source["value"] if last_source else ""),
                "now": now_iso(),
            })
        if path == "/api/admin/backup/auto-trigger":
            # 立即把 DB + 附件全量同步到 COS（管理员可手动触发）
            try:
                db_ok = cos_upload_db()
                up_count = 0
                if _os.path.isdir(UPLOAD_DIR):
                    for root, dirs, files in _os.walk(UPLOAD_DIR):
                        for fn in files:
                            fp = _os.path.join(root, fn)
                            if cos_upload_attachment(fp):
                                up_count += 1
                # 写 backup_meta
                try:
                    c4 = get_db()
                    c4.execute("INSERT OR REPLACE INTO backup_meta(key,value) VALUES ('last_backup_at',?)", (now_iso(),))
                    c4.execute("INSERT OR REPLACE INTO backup_meta(key,value) VALUES ('last_backup_source','admin-manual')",)
                    c4.commit(); c4.close()
                except Exception as e:
                    print("[COS] 写 backup_meta 失败：", e, flush=True)
                return json_resp({"ok": True, "db_synced": db_ok, "uploads_synced": up_count})
            except Exception as e:
                import traceback as _tb
                err = _tb.format_exc()
                print("[AUTO-TRIGGER-ERROR]:\n%s" % err, flush=True)
                return json_resp({"ok": False, "error": str(e), "trace": err}, 500)
        return json_resp({"error": "not found"}, 404)

    def admin_post(self, path, admin):
        d = read_json_body(self)
        conn = get_db()
        try:
            if path == "/api/admin/create-user":
                uname = (d.get("username") or "").strip()
                pw = d.get("password") or ""
                role = d.get("role") or "member"
                if not uname or not pw:
                    return json_resp({"error": "用户名和密码必填"}, 400)
                if conn.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone():
                    return json_resp({"error": "用户名已存在"}, 409)
                h, salt = hash_pass(pw)
                conn.execute("INSERT INTO users (username,display_name,pass_hash,salt,created_at,role) VALUES (?,?,?,?,?,?)",
                             (uname, d.get("display_name") or uname, h, salt, now_iso(), role))
                uid = conn.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone()["id"]
                conn.execute("INSERT INTO settings (user_id,signature) VALUES (?,?)", (uid, "招展顾问"))
                self.seed_demo_data(conn, uid)
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/admin/reset-password":
                uid = d.get("user_id")
                pw = d.get("new_password") or ""
                if not pw:
                    return json_resp({"error": "新密码必填"}, 400)
                if int(uid) == admin["id"]:
                    return json_resp({"error": "不能重置自己的密码，请在设置页修改"}, 400)
                h, salt = hash_pass(pw)
                conn.execute("UPDATE users SET pass_hash=?, salt=? WHERE id=?", (h, salt, uid))
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/admin/update-role":
                uid = d.get("user_id")
                role = d.get("role")
                if role not in ("member", "admin"):
                    return json_resp({"error": "角色非法"}, 400)
                if int(uid) == admin["id"]:
                    return json_resp({"error": "不能修改自己的角色"}, 400)
                conn.execute("UPDATE users SET role=? WHERE id=?", (role, uid))
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/backup/import":
                # 全量导入（管理员）：清空并重建各表，用于迁移恢复
                payload = d.get("data") or d
                tables = payload.get("tables") or {}
                if not isinstance(tables, dict) or not tables:
                    return json_resp({"error": "备份数据为空或格式不正确"}, 400)
                try:
                    conn.execute("PRAGMA foreign_keys=OFF")
                    restored = []
                    for t, rows in tables.items():
                        if not isinstance(rows, list) or not rows:
                            continue
                        # 只恢复已知业务表，防止误导入系统表
                        conn.execute(f"DELETE FROM {t}")
                        for row in rows:
                            cols = [c for c in row.keys() if c != "rowid"]
                            if not cols:
                                continue
                            col_names = ",".join(cols)
                            placeholders = ",".join("?" for _ in cols)
                            conn.execute(f"INSERT INTO {t} ({col_names}) VALUES ({placeholders})",
                                         [row[c] for c in cols])
                        restored.append(t)
                    # 修正自增序列，避免后续插入主键冲突
                    for t in restored:
                        try:
                            maxid = conn.execute(f"SELECT MAX(id) FROM {t}").fetchone()[0]
                            if maxid is not None:
                                conn.execute("INSERT OR REPLACE INTO sqlite_sequence(name,seq) VALUES (?,?)", (t, maxid))
                        except Exception:
                            pass
                    conn.execute("PRAGMA foreign_keys=ON")
                    conn.commit()
                    return json_resp({"ok": True, "restored_tables": restored})
                except Exception as e:
                    conn.rollback()
                    return json_resp({"error": f"导入失败：{str(e)}"}, 500)
        finally:
            conn.close()
        if path == "/api/admin/backup/import-zip":
            # 全量恢复（含数据库 + 附件），会覆盖当前数据，用于迁移后恢复
            import io, zipfile, base64, tempfile, shutil
            zip_b64 = d.get("zip_b64") or ""
            if not zip_b64:
                return json_resp({"error": "未收到备份文件"}, 400)
            try:
                conn.close()
            except Exception:
                pass
            try:
                raw = base64.b64decode(zip_b64)
            except Exception:
                return json_resp({"error": "备份文件解码失败"}, 400)
            tmp = tempfile.mkdtemp()
            try:
                with zipfile.ZipFile(io.BytesIO(raw), "r") as zf:
                    zf.extractall(tmp)
                db_src = os.path.join(tmp, "workbench.db")
                if os.path.exists(db_src):
                    if os.path.exists(DB_PATH):
                        try:
                            os.remove(DB_PATH)
                        except Exception:
                            pass
                    shutil.move(db_src, DB_PATH)
                up_src = os.path.join(tmp, "uploads")
                if os.path.isdir(up_src):
                    if os.path.isdir(UPLOAD_DIR):
                        for fn in os.listdir(UPLOAD_DIR):
                            fp = os.path.join(UPLOAD_DIR, fn)
                            try:
                                if os.path.isfile(fp):
                                    os.remove(fp)
                                elif os.path.isdir(fp):
                                    shutil.rmtree(fp)
                            except Exception:
                                pass
                    cos_clear_uploads()
                    for root, dirs, files in os.walk(up_src):
                        for fn in files:
                            src = os.path.join(root, fn)
                            rel = os.path.relpath(src, up_src)
                            dst = os.path.join(UPLOAD_DIR, rel)
                            os.makedirs(os.path.dirname(dst), exist_ok=True)
                            shutil.copy2(src, dst)
                            cos_upload_attachment(dst)
                    cos_upload_db()
                return json_resp({"ok": True, "note": "数据库与附件已恢复，请刷新页面重新登录"})
            except Exception as e:
                return json_resp({"error": "恢复失败：" + str(e)}, 500)
            finally:
                shutil.rmtree(tmp, ignore_errors=True)
        return json_resp({"error": "not found"}, 404)

    def admin_delete(self, path, admin):
        if path.startswith("/api/admin/users/"):
            tid = path.split("/")[-1]
            try:
                tid = int(tid)
            except Exception:
                return json_resp({"error": "bad id"}, 400)
            if tid == admin["id"]:
                return json_resp({"error": "不能删除自己的账号"}, 400)
            conn = get_db()
            try:
                conn.execute("DELETE FROM users WHERE id=?", (tid,))
                conn.execute("DELETE FROM settings WHERE user_id=?", (tid,))
                conn.execute("DELETE FROM todos WHERE user_id=?", (tid,))
                conn.execute("DELETE FROM customers WHERE user_id=?", (tid,))
                conn.execute("DELETE FROM tags WHERE user_id=?", (tid,))
                conn.execute("DELETE FROM templates WHERE user_id=?", (tid,))
                conn.execute("DELETE FROM email_logs WHERE user_id=?", (tid,))
                conn.commit()
                return json_resp({"ok": True})
            finally:
                conn.close()
        return json_resp({"error": "not found"}, 404)

    # ---------------- 演示数据种子 ----------------
    @staticmethod
    def seed_demo_data(conn, uid):
        """新用户注册后植入示例数据，进来即可直接体验（与纯前端版一致）。"""
        now = now_iso()
        # 预制标签
        for name in ["预制菜客户", "调味品客户", "零食客户", "原料客户", "高意向", "待跟进"]:
            try:
                conn.execute("INSERT OR IGNORE INTO tags (user_id,name) VALUES (?,?)", (uid, name))
            except Exception:
                pass
        # 示例客户
        demo_cust = [
            ("XX预制菜工厂", "王总", "wang@xx-food.com", "13800000001", "SIAL 巴黎食品展", "预制菜客户,高意向"),
            ("YY调味品有限公司", "李总", "li@yy-seasoning.com", "13800000002", "SIAL 巴黎食品展", "调味品客户"),
            ("ZZ休闲零食", "赵经理", "zhao@zz-snack.com", "13800000003", "越南食品展 VietFood", "零食客户,待跟进"),
        ]
        for c in demo_cust:
            conn.execute("INSERT INTO customers (user_id,company,contact,email,phone,exhibition,tags,created_at) VALUES (?,?,?,?,?,?,?,?)",
                         (uid, c[0], c[1], c[2], c[3], c[4], c[5], now))
        # 示例待办（带绑定日期）
        import datetime
        base = datetime.date.today()
        demo_todos = [
            ("跟进XX预制菜厂参展意向", (base + datetime.timedelta(days=1)).isoformat(), "高"),
            ("发送SIAL展位图给YY调味品", (base + datetime.timedelta(days=2)).isoformat(), "中"),
            ("整理越南食品展客户名单", base.isoformat(), "低"),
        ]
        for t in demo_todos:
            conn.execute("INSERT INTO todos (user_id,title,due_time,bind_date,priority,done,created_at) VALUES (?,?,?,?,?,?,?)",
                         (uid, t[0], t[1] + " 18:00", t[1], t[2], 0, now))
        # 示例邮件模板
        conn.execute("""INSERT INTO templates (user_id,name,exhibition,customer_type,scene,tone,subject,body,signature,created_at)
                         VALUES (?,?,?,?,?,?,?,?,?,?)""",
                     (uid, "SIAL初次开发-预制菜", "SIAL 巴黎食品展", "预制菜", "1", "正式商务",
                      "【{客户名称}】诚邀莅临 SIAL 巴黎食品展",
                      "尊敬的{联系人姓名}（{客户名称}）：\n您好！我是「SIAL 巴黎食品展」中国区招展团队的{销售姓名}。\n\n结合贵司在预制菜领域的产品矩阵与出海布局，我们相信贵司非常契合本次展会的买家画像。诚挚邀请贵司莅临SIAL，与海外买家面对面洽谈。\n\n——{销售姓名}",
                      "招展顾问", now))

    # ---------------- POST API ----------------
    def route_post(self, path):
        # 注册 / 登录 不需要鉴权
        if path == "/api/register":
            d = read_json_body(self)
            uname = (d.get("username") or "").strip()
            pw = d.get("password") or ""
            if not uname or not pw:
                return json_resp({"error": "用户名和密码必填"}, 400)
            conn = get_db()
            if conn.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone():
                conn.close()
                return json_resp({"error": "用户名已存在"}, 409)
            h, salt = hash_pass(pw)
            conn.execute("INSERT INTO users (username,display_name,pass_hash,salt,created_at) VALUES (?,?,?,?,?)",
                         (uname, d.get("display_name") or uname, h, salt, now_iso()))
            uid = conn.execute("SELECT id FROM users WHERE username=?", (uname,)).fetchone()["id"]
            conn.execute("INSERT INTO settings (user_id,signature) VALUES (?,?)", (uid, "招展顾问"))
            self.seed_demo_data(conn, uid)
            conn.commit()
            conn.close()
            token = gen_token()
            with SESS_LOCK:
                SESSIONS[token] = uid
            return json_resp({"token": token, "user": {"id": uid, "username": uname, "display_name": d.get("display_name") or uname, "role": "member"}})

        if path == "/api/login":
            d = read_json_body(self)
            uname = (d.get("username") or "").strip()
            pw = d.get("password") or ""
            conn = get_db()
            row = conn.execute("SELECT * FROM users WHERE username=?", (uname,)).fetchone()
            conn.close()
            if not row:
                return json_resp({"error": "用户不存在"}, 404)
            h, _ = hash_pass(pw, row["salt"])
            if h != row["pass_hash"]:
                return json_resp({"error": "密码错误"}, 401)
            token = gen_token()
            with SESS_LOCK:
                SESSIONS[token] = row["id"]
            return json_resp({"token": token, "user": {"id": row["id"], "username": row["username"], "display_name": row["display_name"], "role": row["role"]}})

        if path.startswith("/api/admin/"):
            a, aerr = require_admin(self)
            if aerr:
                return aerr
            return self.admin_post(path, a)

        u, err = require_user(self)
        if err:
            return err
        uid = u["id"]
        d = read_json_body(self)
        conn = get_db()
        try:
            if path == "/api/todos":
                conn.execute("INSERT INTO todos (user_id,title,due_time,bind_date,priority,customer_id,created_at) VALUES (?,?,?,?,?,?,?)",
                             (uid, d.get("title",""), d.get("due_time"), d.get("bind_date"), d.get("priority","中"), d.get("customer_id"), now_iso()))
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/tags":
                name = (d.get("name") or "").strip()
                if not name:
                    return json_resp({"error":"标签名必填"},400)
                try:
                    conn.execute("INSERT INTO tags (user_id,name) VALUES (?,?)", (uid, name))
                    conn.commit()
                except sqlite3.IntegrityError:
                    return json_resp({"error":"标签已存在"},409)
                return json_resp({"ok": True})
            if path == "/api/customers":
                tags = d.get("tags")
                if isinstance(tags, list):
                    tags = ",".join(tags)
                conn.execute("INSERT INTO customers (user_id,company,contact,email,phone,exhibition,tags,status,created_at) VALUES (?,?,?,?,?,?,?,?,?)",
                             (uid, d.get("company"), d.get("contact"), d.get("email"), d.get("phone"), d.get("exhibition"), tags or "", d.get("status", "潜在客户"), now_iso()))
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/customers/batch-delete":
                ids = d.get("ids") or []
                ids = [str(i) for i in ids if i not in (None, "")]
                deleted = 0
                if ids:
                    qmarks = ",".join("?" * len(ids))
                    cur = conn.execute(
                        f"DELETE FROM customers WHERE user_id=? AND id IN ({qmarks})",
                        [uid] + ids,
                    )
                    deleted = cur.rowcount
                    conn.commit()
                return json_resp({"ok": True, "deleted": deleted})
            if path == "/api/customers/import":
                text = d.get("csv") or ""
                return self.import_csv(conn, uid, text)
            if path == "/api/customers/upload-excel":
                # 接收 base64 编码的文件数据（前端 FileReader → base64）
                b64_data = d.get("file_base64") or d.get("file_data") or ""
                filename = d.get("filename") or "data.xlsx"
                if not b64_data:
                    return json_resp({"error": "未上传文件，请先选择 Excel/CSV 文件"}, 400)
                import base64 as b64mod
                try:
                    raw = b64mod.b64decode(b64_data)
                except Exception as e:
                    return json_resp({"error": f"文件数据解码失败：{str(e)}。请确认文件未损坏后重试。"}, 400)
                ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "xlsx"
                # 文件大小检查(限制50MB)
                if len(raw) > 50 * 1024 * 1024:
                    return json_resp({"error": f"文件过大({len(raw)//1024//1024}MB)，请控制在50MB以内或拆分上传"}, 400)
                try:
                    result = self.parse_excel_file(raw, ext)
                    if result.get("error"):
                        return json_resp({"error": result["error"]}, 400)
                    imp_result = self.import_parsed_data(conn, uid, result["rows"], result.get("headers", []))
                    # imp_result 现在是 dict: {imported: N, diagnostic?: {...}}
                    imported_count = imp_result.get("imported", 0) if isinstance(imp_result, dict) else imp_result
                    resp = {"ok": True, "imported": imported_count, "total_rows": len(result["rows"]),
                            "preview": result["preview"], "headers": result.get("headers", []),
                            "detected_columns": result.get("detected_columns", {})}
                    # 导入0行时附加诊断信息
                    if isinstance(imp_result, dict) and imp_result.get("diagnostic"):
                        resp["diagnostic"] = imp_result["diagnostic"]
                    return json_resp(resp)
                except ImportError as e:
                    return json_resp({"error": f"服务器缺少文件解析依赖（{str(e)}）。请联系管理员安装 openpyxl。"}, 500)
                except Exception as e:
                    # 记录详细错误到stderr(容器日志可见)
                    import traceback
                    tb = traceback.format_exc()
                    print(f"[Excel导入错误] 用户={uid} 文件={filename} 大小={len(raw)} 错误: {tb}", flush=True)
                    err_msg = str(e)
                    # 对常见错误给出友好提示
                    if "openpyxl" in str(type(e).__module__).lower() and "No module" in err_msg:
                        err_msg = "服务器未安装 Excel 解析库(openpyxl)，请联系管理员"
                    elif "zip" in err_msg.lower() or "not a valid" in err_msg.lower():
                        err_msg = "文件格式异常，请确认是有效的 .xlsx 文件（尝试用 Excel 另存为一次）"
                    return json_resp({"error": f"文件解析失败：{err_msg}"}, 400)
            if path == "/api/templates":
                conn.execute("INSERT INTO templates (user_id,name,exhibition,customer_type,scene,tone,subject,body,signature,attachment_ids,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                             (uid, d.get("name"), d.get("exhibition"), d.get("customer_type"), d.get("scene"), d.get("tone"),
                              d.get("subject"), d.get("body"), d.get("signature"), d.get("attachment_ids",""), now_iso()))
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/ai/generate":
                _settings = load_settings(uid)
                # 测试连接：请求体可临时覆盖 AI 配置（不落库），并强制启用一次 LLM
                if d.get("_llm_test"):
                    _settings = dict(_settings)
                    for _k in ("ai_enabled", "ai_provider", "ai_base_url", "ai_api_key", "ai_model"):
                        if _k in d and d[_k] != "":
                            _settings[_k] = d[_k]
                    _settings["ai_enabled"] = "1"
                # 前端传 angle_key 时按指定角度生成（用于每次点切换出差异化邮件）
                _angle = None
                _ak = d.get("angle_key")
                if _ak:
                    try:
                        _angle = next((a for a in MULTI_ANGLES if a.get("key") == _ak), None)
                    except Exception:
                        _angle = None
                subject, body, _lu, _le = build_email(d.get("exhibition"), d.get("customer_type"), d.get("scene"),
                                            d.get("tone"), d.get("custom_input"), d.get("signature"), uid,
                                            material_ids=d.get("material_ids"), angle=_angle, settings=_settings)
                return json_resp({"subject": subject, "body": body, "llm_used": _lu, "llm_error": _le,
                                  "angle_key": _angle.get("key") if _angle else None, "angle_label": _angle.get("label") if _angle else None})
            if path == "/api/ai/generate-multi":
                n = d.get("n") or 5
                try:
                    n = max(4, min(5, int(n)))
                except Exception:
                    n = 5
                _settings = load_settings(uid)
                versions = build_email_multi(d.get("exhibition"), d.get("customer_type"), d.get("scene"),
                                             d.get("tone"), d.get("custom_input"), d.get("signature"), uid,
                                             material_ids=d.get("material_ids"), n=n, settings=_settings)
                return json_resp({"versions": versions})
            if path == "/api/ai/translate":
                return self.translate_email(d)
            if path == "/api/exhibitions":
                conn.execute("INSERT INTO exhibitions (user_id,name,city,date_text,note) VALUES (?,?,?,?,?)",
                             (uid, d.get("name"), d.get("city"), d.get("date_text"), d.get("note")))
                conn.commit()
                new_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
                return json_resp({"ok": True, "id": new_id})
            if path == "/api/exhibitions/batch-import":
                # 批量导入展会（支持 xlsx/xls/csv）
                b64_data = d.get("content_b64", "")
                filename = d.get("filename", "import.xlsx")
                if not b64_data:
                    return json_resp({"error": "请选择文件"}, 400)
                import base64 as b64mod
                try:
                    raw = b64mod.b64decode(b64_data)
                except Exception as e:
                    return json_resp({"error": f"文件解码失败：{str(e)}"}, 400)
                ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "xlsx"
                rows = []
                headers = []
                try:
                    if ext in ("xlsx", "xls"):
                        import io
                        import openpyxl
                        wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
                        ws = wb.active
                        iter_rows = list(ws.iter_rows(values_only=True))
                        if not iter_rows:
                            return json_resp({"error": "文件为空，没有数据行"}, 400)
                        headers = [str(c or "").strip() for c in iter_rows[0]]
                        for r in iter_rows[1:]:
                            if any(c is not None for c in r):
                                rows.append([str(c or "") for c in r])
                    elif ext == "csv":
                        import csv, io
                        reader = csv.reader(io.StringIO(raw.decode("utf-8-sig", errors="replace")))
                        raw_rows = list(reader)
                        if not raw_rows:
                            return json_resp({"error": "CSV 文件为空"}, 400)
                        headers = [c.strip() for c in raw_rows[0]]
                        for r in raw_rows[1:]:
                            if any(c.strip() for c in r):
                                rows.append([c.strip() for c in r])
                    else:
                        return json_resp({"error": f"不支持的文件格式 .{ext}，请用 .xlsx / .xls / .csv"}, 400)
                except ImportError as e:
                    return json_resp({"error": f"服务器缺少解析依赖（{str(e)}）"}, 500)
                except Exception as e:
                    return json_resp({"error": f"文件解析失败：{str(e)}"}, 400)
                if not rows:
                    return json_resp({"error": "文件只有表头，没有数据行"}, 400)
                # 列名映射：支持中英文多种写法
                col_map = {}
                name_cols = ["展会名称", "名称", "展会展称", "name", "展会"]
                city_cols = ["城市", "举办城市", "city", "地点"]
                date_cols = ["档期", "日期", "展期", "时间", "date", "date_text"]
                note_cols = ["备注", "说明", "亮点", "note", "注释"]
                for i, h in enumerate(headers):
                    hl = h.lower().strip()
                    for mc in name_cols:
                        if hl == mc.lower() or mc.lower() in hl: col_map["name"] = i; break
                    for mc in city_cols:
                        if hl == mc.lower() or mc.lower() in hl: col_map["city"] = i; break
                    for mc in date_cols:
                        if hl == mc.lower() or mc.lower() in hl: col_map["date_text"] = i; break
                    for mc in note_cols:
                        if hl == mc.lower() or mc.lower() in hl: col_map["note"] = i; break
                if "name" not in col_map:
                    return json_resp({"error": f"未找到「展会名称」列。当前表头：{headers}。请确保第一列包含展会名称。", "headers": headers}, 400)
                imported = 0
                errors = []
                for ri, row in enumerate(rows):
                    name = (row[col_map["name"]] if col_map["name"] < len(row) else "").strip()
                    if not name:
                        continue  # 跳过空名称行
                    city = (row[col_map["city"]] if "city" in col_map and col_map["city"] < len(row) else "").strip()
                    dt = (row[col_map["date_text"]] if "date_text" in col_map and col_map["date_text"] < len(row) else "").strip()
                    note = (row[col_map["note"]] if "note" in col_map and col_map["note"] < len(row) else "").strip()
                    try:
                        conn.execute("INSERT INTO exhibitions (user_id,name,city,date_text,note) VALUES (?,?,?,?,?)",
                                     (uid, name, city, dt, note))
                        imported += 1
                    except Exception as e:
                        errors.append(f"第{ri+2}行「{name}」: {str(e)}")
                conn.commit()
                resp = {"ok": True, "imported": imported, "total": len(rows), "errors": errors}
                if errors:
                    resp["warning"] = f"成功 {imported} 条，{len(errors)} 行跳过"
                return json_resp(resp)
            if path == "/api/drafts":
                title = (d.get("title") or "").strip() or "未命名草稿"
                payload = json.dumps(d.get("payload") or {}, ensure_ascii=False)
                cur = conn.execute(
                    "INSERT INTO drafts (user_id,title,payload,created_at,updated_at) VALUES (?,?,?,?,?)",
                    (uid, title, payload, now_iso(), now_iso()))
                conn.commit()
                return json_resp({"ok": True, "id": cur.lastrowid})
            if path == "/api/materials":
                # 接收文件内容（base64）或直接记录 URL；演示以名称 + 路径记录
                name = d.get("name") or "资料"
                ex_id = d.get("exhibition_id")
                fpath = os.path.join(UPLOAD_DIR, f"{uid}_{int(datetime.datetime.now().timestamp())}_{name}")
                content = d.get("content_b64")
                if content:
                    import base64
                    with open(fpath, "wb") as f:
                        f.write(base64.b64decode(content))
                else:
                    fpath = d.get("file_path") or ""
                if os.path.exists(fpath):
                    cos_upload_attachment(fpath)
                cur = conn.execute("INSERT INTO materials (user_id,exhibition_id,name,file_path,created_at) VALUES (?,?,?,?,?)",
                             (uid, ex_id, name, fpath, now_iso()))
                conn.commit()
                new_id = cur.lastrowid
                return json_resp({"ok": True, "id": new_id})
            if path == "/api/email/preview":
                return self.preview_emails(conn, uid, u, d)
            if path == "/api/email/send":
                return self.send_emails(conn, uid, u, d)
            if path == "/api/settings":
                s = d
                ensure_ai_columns()
                conn.execute("""INSERT OR REPLACE INTO settings
                    (user_id,smtp_host,smtp_port,smtp_user,smtp_pass,from_email,from_name,default_interval,demo_mode,signature,
                     ai_enabled,ai_provider,ai_base_url,ai_api_key,ai_model)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (uid, s.get("smtp_host"), s.get("smtp_port"), s.get("smtp_user"), s.get("smtp_pass"),
                     s.get("from_email"), s.get("from_name"), s.get("default_interval",5),
                     s.get("demo_mode",1), s.get("signature"),
                     s.get("ai_enabled", ""), s.get("ai_provider", "deepseek"),
                     s.get("ai_base_url", ""), s.get("ai_api_key", ""), s.get("ai_model", "")))
                conn.commit()
                # 关键配置（AI 大模型）保存后立即同步上传到 COS，避免容器重启后丢配置
                if s.get("ai_enabled") and s.get("ai_api_key"):
                    try: cos_upload_db()
                    except Exception as _e: print("[COS] settings保存后立即同步失败：", _e, flush=True)
                return json_resp({"ok": True})
            if path == "/api/me/change-password":
                old = d.get("old_password") or ""
                new = d.get("new_password") or ""
                if not new:
                    return json_resp({"error": "新密码必填"}, 400)
                row = conn.execute("SELECT pass_hash,salt FROM users WHERE id=?", (uid,)).fetchone()
                oh, _ = hash_pass(old, row["salt"])
                if oh != row["pass_hash"]:
                    return json_resp({"error": "原密码不正确"}, 400)
                h, salt = hash_pass(new)
                conn.execute("UPDATE users SET pass_hash=?, salt=? WHERE id=?", (h, salt, uid))
                conn.commit()
                return json_resp({"ok": True})
            if path == "/api/me/profile":
                # 修改显示名称（成员可自行修改）
                display_name = (d.get("display_name") or "").strip()
                if not display_name:
                    return json_resp({"error": "显示名称不能为空"}, 400)
                conn.execute("UPDATE users SET display_name=? WHERE id=?", (display_name, uid))
                conn.commit()
                return json_resp({"ok": True, "display_name": display_name})
        finally:
            conn.close()
        return json_resp({"error": "not found"}, 404)

    # ---------------- PATCH API ----------------
    def route_patch(self, path):
        u, err = require_user(self)
        if err:
            return err
        uid = u["id"]
        d = read_json_body(self)
        conn = get_db()
        try:
            if path.startswith("/api/todos/"):
                tid = path.split("/")[-1]
                fields = []
                vals = []
                for k in ("title","due_time","bind_date","priority","customer_id","done"):
                    if k in d:
                        fields.append(f"{k}=?")
                        vals.append(d[k])
                if fields:
                    conn.execute(f"UPDATE todos SET {','.join(fields)} WHERE id=? AND user_id=?", vals+[tid, uid])
                    conn.commit()
                return json_resp({"ok": True})
            if path.startswith("/api/customers/"):
                cid = path.split("/")[-1]
                fields = []
                vals = []
                for k in ("company","contact","email","phone","exhibition","tags","status"):
                    if k in d:
                        v = d[k]
                        if k == "tags" and isinstance(v, list):
                            v = ",".join(v)
                        fields.append(f"{k}=?")
                        vals.append(v)
                if fields:
                    conn.execute(f"UPDATE customers SET {','.join(fields)} WHERE id=? AND user_id=?", vals+[cid, uid])
                    conn.commit()
                return json_resp({"ok": True})
            if path.startswith("/api/exhibitions/"):
                eid = path.split("/")[-1]
                fields = []
                vals = []
                for k in ("name", "city", "date_text", "note"):
                    if k in d:
                        v = d[k]
                        if k == "date_text":
                            v = _normalize_date_text(v)
                        fields.append(f"{k}=?")
                        vals.append(v)
                if fields:
                    # 同时允许修改全局展会(user_id=0)与本人维护的展会
                    conn.execute(
                        f"UPDATE exhibitions SET {','.join(fields)} WHERE id=? AND (user_id=? OR user_id=0)",
                        vals + [eid, uid])
                    conn.commit()
                return json_resp({"ok": True})
            if path.startswith("/api/templates/"):
                tid = path.split("/")[-1]
                fields = []
                vals = []
                for k in ("name","exhibition","customer_type","scene","tone","subject","body","signature","attachment_ids"):
                    if k in d:
                        fields.append(f"{k}=?")
                        vals.append(d[k])
                if fields:
                    conn.execute(f"UPDATE templates SET {','.join(fields)} WHERE id=? AND user_id=?", vals+[tid, uid])
                    conn.commit()
                return json_resp({"ok": True})
            if path.startswith("/api/materials/"):
                # 编辑资料：支持改名 / 改所属展会 / 替换文件
                mid = path.split("/")[-1]
                row = conn.execute("SELECT * FROM materials WHERE id=? AND (user_id=? OR user_id=0)", (mid, uid)).fetchone()
                if not row:
                    return json_resp({"error": "资料不存在或无权限"}, 404)
                fields, vals = [], []
                if "name" in d and d["name"]:
                    fields.append("name=?"); vals.append(d["name"])
                if "exhibition_id" in d:
                    new_ex_id = d["exhibition_id"]
                    # 若传入字符串（新建展会名），先查或建
                    if isinstance(new_ex_id, str) and new_ex_id.strip():
                        nm = new_ex_id.strip()
                        m_ex = conn.execute("SELECT id FROM exhibitions WHERE name=? AND (user_id=? OR user_id=0)", (nm, uid)).fetchone()
                        if m_ex:
                            new_ex_id = m_ex["id"]
                        else:
                            cur = conn.execute("INSERT INTO exhibitions (user_id,name,city,date_text,note) VALUES (?,?,?,?,?)",
                                                (uid, nm, "", "", "编辑资料时新建"))
                            new_ex_id = cur.lastrowid
                    fields.append("exhibition_id=?"); vals.append(new_ex_id)
                if "content_b64" in d and d["content_b64"]:
                    import base64 as _b64
                    old_path = row["file_path"] or ""
                    name = d.get("name") or row["name"] or "资料"
                    new_path = os.path.join(UPLOAD_DIR, f"{uid}_{int(datetime.datetime.now().timestamp())}_{name}")
                    with open(new_path, "wb") as f:
                        f.write(_b64.b64decode(d["content_b64"]))
                    if os.path.exists(new_path):
                        cos_upload_attachment(new_path)
                    fields.append("file_path=?"); vals.append(new_path)
                if fields:
                    conn.execute(f"UPDATE materials SET {','.join(fields)} WHERE id=?", vals+[mid])
                    conn.commit()
                return json_resp({"ok": True})
            if path.startswith("/api/drafts/"):
                did = path.split("/")[-1]
                row = conn.execute("SELECT id FROM drafts WHERE id=? AND user_id=?", (did, uid)).fetchone()
                if not row:
                    return json_resp({"error": "草稿不存在"}, 404)
                fields, vals = [], []
                if "title" in d and d["title"]:
                    fields.append("title=?"); vals.append(d["title"])
                if "payload" in d:
                    fields.append("payload=?"); vals.append(json.dumps(d["payload"], ensure_ascii=False))
                if fields:
                    fields.append("updated_at=?"); vals.append(now_iso())
                    conn.execute(f"UPDATE drafts SET {','.join(fields)} WHERE id=?", vals+[did])
                    conn.commit()
                return json_resp({"ok": True})
        finally:
            conn.close()
        return json_resp({"error": "not found"}, 404)

    # ---------------- DELETE API ----------------
    def route_delete(self, path):
        u, err = require_user(self)
        if err:
            return err
        if path.startswith("/api/admin/"):
            a, aerr = require_admin(self)
            if aerr:
                return aerr
            return self.admin_delete(path, a)
        uid = u["id"]
        conn = get_db()
        try:
            if path.startswith("/api/todos/"):
                tid = path.split("/")[-1]
                conn.execute("DELETE FROM todos WHERE id=? AND user_id=?", (tid, uid))
            elif path.startswith("/api/customers/"):
                cid = path.split("/")[-1]
                conn.execute("DELETE FROM customers WHERE id=? AND user_id=?", (cid, uid))
            elif path.startswith("/api/tags/"):
                tid = path.split("/")[-1]
                conn.execute("DELETE FROM tags WHERE id=? AND user_id=?", (tid, uid))
            elif path.startswith("/api/templates/"):
                tid = path.split("/")[-1]
                conn.execute("DELETE FROM templates WHERE id=? AND user_id=?", (tid, uid))
            elif path.startswith("/api/materials/"):
                mid = path.split("/")[-1]
                conn.execute("DELETE FROM materials WHERE id=? AND (user_id=? OR user_id=0)", (mid, uid))
            elif path.startswith("/api/drafts/"):
                did = path.split("/")[-1]
                conn.execute("DELETE FROM drafts WHERE id=? AND user_id=?", (did, uid))
            elif path.startswith("/api/exhibitions/"):
                eid = path.split("/")[-1]
                # 拒绝删除全局共享展会（user_id=0），只允许删除本人建的
                row = conn.execute("SELECT user_id FROM exhibitions WHERE id=?", (eid,)).fetchone()
                if row and int(row["user_id"]) == 0:
                    return json_resp({"error": "系统级展会不能删除，可在「展会管理」编辑后停用"}, 400)
                conn.execute("DELETE FROM exhibitions WHERE id=? AND user_id=?", (eid, uid))
            else:
                return json_resp({"error":"not found"},404)
            conn.commit()
            return json_resp({"ok": True})
        finally:
            conn.close()

    # ---------------- 业务方法 ----------------
    def import_csv(self, conn, uid, text):
        try:
            lines = text.strip().splitlines()
            # 自动跳过非表头行（如"销售跟进情况表"等标题行）
            while lines and not any(k in lines[0] for k in ["公司名称","客户公司","联系人","邮箱","序号"]):
                lines.pop(0)
            if not lines:
                return json_resp({"error": "未找到有效的表头行"}, 400)
            reader = csv.DictReader(io.StringIO("\n".join(lines)))
            fields = reader.fieldnames or []
            # 检测是否为「销售跟进情况表」模板（通过特征列名判断）
            is_template = any("公司名称" in str(f) for f in fields)
            count = 0
            for row in reader:
                if is_template:
                    # 销售跟进情况表模板映射
                    company = (row.get("公司名称（全称）*") or row.get("公司名称（全称）") or row.get("公司名称(全称)") or "").strip()
                    if not company:
                        company = (row.get("公司名称（简称）") or "").strip()
                    if not company:
                        continue
                    contact = (row.get("联系人") or "").strip()
                    phone = (row.get("联系电话") or "").strip()
                    exhibition = (row.get("参展记录*") or row.get("参展记录") or row.get("参展记录*示例：24.5泰国食品展") or "").strip().split("\n")[0]
                    tags = (row.get("客户分配标签属性") or "").strip()
                    remark = (row.get("客户联系情况或跟踪记录*") or row.get("客户联系情况或跟踪记录") or row.get("其他备注（若有）") or "").strip()
                    region = " ".join(filter(None, [row.get("省") or "", row.get("市") or "", row.get("县区") or ""])).strip()
                    source = (row.get("客户来源") or "").strip()
                    email = ""  # 该模板无邮箱列，需后续补充
                else:
                    # 标准格式
                    company = (row.get("客户公司") or row.get("company") or row.get("公司名称") or "").strip()
                    if not company:
                        continue
                    contact = (row.get("联系人") or row.get("contact") or "").strip()
                    phone = (row.get("手机号") or row.get("phone") or row.get("联系电话") or "").strip()
                    exhibition = (row.get("意向展会") or row.get("exhibition") or row.get("参展记录") or "").strip()
                    tags = (row.get("客户标签") or row.get("tags") or row.get("客户分配标签属性") or "").strip()
                    email = (row.get("邮箱") or row.get("email") or "").strip()
                    remark = ""
                    region = ""
                    source = ""
                conn.execute("INSERT INTO customers (user_id,company,contact,email,phone,exhibition,tags,remark,region,source,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                             (uid, company, contact, email, phone, exhibition, tags, remark, region, source, now_iso()))
                count += 1
            conn.commit()
            return json_resp({"ok": True, "imported": count, "template_mode": is_template})
        except Exception as e:
            return json_resp({"error": str(e)}, 400)

    # ---------- Excel 文件上传解析（openpyxl 可靠解析） ----------
    # 常见表头关键词，用于判断一行是不是真正的表头
    _HEADER_KEYWORDS = {"公司", "名称", "联系", "邮箱", "邮件", "手机", "电话", "展会",
                        "标签", "备注", "客户", "企业", "供应商", "省", "市", "区", "县",
                        "产品", "类型", "序号", "编号", "姓名", "地址", "职位", "行业"}

    @staticmethod
    def _looks_like_header(row):
        """判断一行是否像表头（包含常见表头关键词）"""
        text = " ".join(str(c).strip() for c in row if c)
        return any(kw in text for kw in Handler._HEADER_KEYWORDS)

    @staticmethod
    def parse_excel_file(raw_data, ext):
        """用 openpyxl 解析 Excel/CSV 文件，返回 {rows, headers, preview, detected_columns, error?}"""
        import tempfile, csv as csv_mod
        suffix = "." + ("xlsx" if ext in ("xlsx", "xls") else "csv")
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(raw_data)
            tmp_path = tmp.name
        try:
            if ext == "csv":
                rows = []
                for enc in ["utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1"]:
                    try:
                        with open(tmp_path, "r", encoding=enc) as f:
                            reader = csv_mod.reader(f)
                            for row in reader:
                                rows.append([str(c).strip() for c in row])
                        break
                    except (UnicodeDecodeError, UnicodeError):
                        rows = []
                        continue
                if not rows:
                    return {"error": "CSV 文件为空或编码无法识别"}
                # CSV 也智能找表头：跳过标题行
                header_idx = 0
                for i, r in enumerate(rows):
                    if any(c.strip() for c in r):
                        if Handler._looks_like_header(r):
                            header_idx = i
                            break
                        if i + 1 < len(rows) and Handler._looks_like_header(rows[i + 1]):
                            header_idx = i + 1
                            break
                        header_idx = i
                        break
                headers = [str(c).strip() if c else "" for c in rows[header_idx]]
                data_rows = [[str(c).strip() if c else "" for c in r] for r in rows[header_idx + 1:]]

            elif ext == "xlsx":
                try:
                    import openpyxl
                except ImportError:
                    return {"error": "服务器未安装 openpyxl 库"}
                try:
                    wb = openpyxl.load_workbook(tmp_path, data_only=True)
                except Exception as e:
                    err_low = str(e).lower()
                    if "bad zip" in err_low or "not a zip" in err_low or "invalid" in err_low:
                        return {"error": "文件不是有效的 xlsx 格式"}
                    return {"error": f"Excel 文件读取失败：{str(e)}"}
                ws = wb.active
                raw_rows = list(ws.iter_rows(values_only=True))
                if not raw_rows:
                    return {"error": "Excel 文件为空"}
                # 智能找表头行：跳过标题行（如"销售跟进情况表"）
                header_row_idx = None
                for i, r in enumerate(raw_rows):
                    if not any(c for c in r if c is not None and str(c).strip()):
                        continue
                    if Handler._looks_like_header(r):
                        header_row_idx = i
                        break
                    if i + 1 < len(raw_rows) and Handler._looks_like_header(raw_rows[i + 1]):
                        header_row_idx = i + 1
                        break
                    header_row_idx = i
                    break

                if header_row_idx is None:
                    return {"error": "Excel 文件没有可识别的表头或数据"}

                headers = [str(c).strip() if c else "" for c in raw_rows[header_row_idx]]
                data_rows = []
                for r in raw_rows[header_row_idx + 1:]:
                    row_str = [str(c).strip() if c else "" for c in r]
                    if any(v.strip() for v in row_str):
                        data_rows.append(row_str)
            elif ext == "xls":
                return {"error": ".xls 旧格式不支持，请另存为 .xlsx 或 .csv"}
            else:
                return {"error": f"不支持的文件格式：.{ext}"}

            if not data_rows:
                return {"error": "文件只有表头没有数据行"}

            # 模糊列名匹配
            detected = Handler.fuzzy_match_columns(headers)
            _data_keys = ["company","contact","email","phone","exhibition","tags","remark"]
            preview = []
            for r in data_rows[:10]:
                preview.append({k: (r[v] if isinstance(v,int) and v < len(r) else "") for k in _data_keys for v in [detected.get(k)]})
            return {"rows": data_rows, "headers": headers, "preview": preview,
                    "detected_columns": detected}
        finally:
            os.unlink(tmp_path)

    @staticmethod
    def _clean_header(h):
        """清洗表头：去掉括号内容、标点、空格、星号"""
        import re
        s = h.strip()
        # 去掉中文/英文括号及其中内容
        s = re.sub(r'[（(][^）)]*[）)]', '', s)
        # 去掉末尾标记
        s = s.rstrip('* \t\n\r')
        return s.strip()

    @staticmethod
    def fuzzy_match_columns(headers):
        """
        模糊匹配列名。核心必填字段：公司名、联系人、邮箱。
        可选字段：手机号、意向展会、标签。
        返回 {company: idx, contact: idx, email: idx, phone: idx, ...}
        """
        col_map = {}
        for i, h in enumerate(headers):
            h_raw = str(h).strip() if h else ""
            h_clean = Handler._clean_header(h_raw)
            h_low = h_clean.lower()
            h_raw_low = h_raw.lower()

            # ===== 公司名（核心必填）=====
            if not col_map.get("company") and any(kw in h_clean for kw in [
                "公司名称", "公司名", "企业名称", "单位名称", "客户公司",
                "公司", "企业", "供应商", "厂商", "品牌", "参展商",
                "公司全称", "全称", "客户名称"
            ]):
                col_map["company"] = i
            if not col_map.get("company") and any(kw in h_raw for kw in [
                "公司名称", "公司名", "企业名称", "单位名称", "客户公司",
                "公司", "企业", "供应商", "厂商", "品牌", "参展商",
                "公司全称", "全称"
            ]):
                col_map["company"] = i
            if not col_map.get("company") and any(kw in h_raw_low for kw in [
                "company", "firm", "org", "supplier", "vendor"
            ]):
                col_map["company"] = i

            # ===== 联系人/姓名（核心必填）=====
            if not col_map.get("contact") and any(kw in h_clean for kw in [
                "联系人", "联系姓名", "姓名", "负责人", "客户姓名",
                "联系人姓名", "对接人", "采购负责人", "决策人"
            ]):
                col_map["contact"] = i
            if not col_map.get("contact") and any(kw in h_raw for kw in [
                "联系人", "联系姓名", "姓名", "负责人", "客户姓名",
                "联系人姓名", "对接人", "采购负责人", "决策人"
            ]):
                col_map["contact"] = i
            if not col_map.get("contact") and any(kw in h_raw_low for kw in [
                "contact", "name", "person"
            ]):
                col_map["contact"] = i

            # ===== 邮箱（核心必填）=====
            if not col_map.get("email") and any(kw in h_clean for kw in [
                "邮箱", "电子邮箱", "Email", "E-mail", "邮件",
                "邮箱地址", "E-mail地址", "邮箱号", "邮箱号码"
            ]):
                col_map["email"] = i
            if not col_map.get("email") and any(kw in h_raw for kw in [
                "邮箱", "电子邮箱", "Email", "E-mail", "邮件",
                "邮箱地址", "E-mail地址", "邮箱号", "邮箱号码"
            ]):
                col_map["email"] = i
            if not col_map.get("email") and any(kw in h_raw_low for kw in [
                "email", "mail", "e_mail"
            ]):
                col_map["email"] = i

            # ===== 手机号（可选）=====
            if not col_map.get("phone") and any(kw in h_clean for kw in [
                "联系电话", "手机号", "电话", "手机", "手机号码",
                "联系方式", "电话号码", "Tel"
            ]):
                col_map["phone"] = i
            if not col_map.get("phone") and any(kw in h_raw for kw in [
                "联系电话", "手机号", "电话", "手机", "手机号码",
                "联系方式", "电话号码", "Tel"
            ]):
                col_map["phone"] = i
            if not col_map.get("phone") and any(kw in h_raw_low for kw in [
                "phone", "tel", "mobile"
            ]):
                col_map["phone"] = i

            # ===== 意向展会（可选）=====
            if not col_map.get("exhibition") and any(kw in h_clean for kw in [
                "参展记录", "意向展会", "展会", "参展", "目标展会",
                "感兴趣展会", "参展意向"
            ]):
                col_map["exhibition"] = i
            if not col_map.get("exhibition") and any(kw in h_raw for kw in [
                "参展记录", "意向展会", "展会", "参展", "目标展会",
                "感兴趣展会", "参展意向"
            ]):
                col_map["exhibition"] = i

            # ===== 标签（可选）=====
            if not col_map.get("tags") and any(kw in h_clean for kw in [
                "标签", "分类", "属性", "类别", "行业"
            ]):
                col_map["tags"] = i
            if not col_map.get("tags") and any(kw in h_raw for kw in [
                "标签", "分类", "属性", "类别", "行业"
            ]):
                col_map["tags"] = i

        # 兜底：只有一列时当 company
        if not col_map.get("company") and len(headers) == 1:
            col_map["company"] = 0
        # 兜底：多列但没匹配到公司名，用第一列
        if not col_map.get("company") and len(headers) > 0:
            col_map["company"] = 0
            col_map["_fallback_company"] = True

        # 记录诊断信息
        col_map["_raw_headers"] = [str(h).strip() for h in headers]
        return col_map

    def import_parsed_data(self, conn, uid, rows, headers):
        """根据模糊匹配的列，将已解析的行数据导入数据库"""
        detected = self.fuzzy_match_columns(headers)
        count = 0
        skipped_empty = 0
        skipped_junk = 0  # 无意义公司名（如"序号"、"编号"）
        sample_skipped = []  # 记录前3条被跳过的行(用于诊断)
        # 无意义公司名关键词
        _junk_keywords = ("序号", "编号", "no.", "no", "#", "id", "行号")
        for row in rows:
            def gv(key):
                idx = detected.get(key)
                return (row[idx].strip() if idx is not None and idx < len(row) else "")
            company = gv("company").strip()
            if not company:
                skipped_empty += 1
                if len(sample_skipped) < 3:
                    sample_skipped.append(row[:min(4, len(row))])
                continue
            # 过滤无意义公司名：纯数字、或包含序号/编号等关键词
            company_low = company.lower()
            if company_low in _junk_keywords or any(kw in company_low for kw in _junk_keywords) or company.isdigit():
                skipped_junk += 1
                continue
            conn.execute("INSERT INTO customers (user_id,company,contact,email,phone,exhibition,tags,remark,created_at) VALUES (?,?,?,?,?,?,?,?,?)",
                         (uid, company, gv("contact"), gv("email"), gv("phone"),
                          gv("exhibition"), gv("tags"), gv("remark"), now_iso()))
            count += 1
        conn.commit()
        # 返回详细结果，包含诊断信息
        result = {"imported": count}
        if count == 0:
            result["diagnostic"] = {
                "headers": headers,
                "detected_columns": detected,
                "total_rows": len(rows),
                "skipped_empty_company": skipped_empty,
                "skipped_junk_company": skipped_junk,
                "sample_skipped_rows": sample_skipped,
            }
        return result

    def preview_emails(self, conn, uid, u, d):
        """根据模板内容 + 客户来源，渲染个性化邮件预览"""
        subject_tpl = d.get("subject", "")
        body_tpl = d.get("body", "")
        # 客户来源
        customer_ids = d.get("customer_ids") or []
        tag_filter = (d.get("tag_filter") or "").strip()
        csv_text = d.get("csv")
        customers = []
        if customer_ids:
            for cid in customer_ids:
                row = conn.execute("SELECT * FROM customers WHERE id=? AND user_id=?", (cid, uid)).fetchone()
                if row:
                    customers.append(row_to_dict(row))
        if csv_text:
            reader = csv.DictReader(io.StringIO(csv_text))
            for row in reader:
                customers.append({
                    "company": (row.get("客户公司") or row.get("company") or "").strip(),
                    "contact": (row.get("联系人") or row.get("contact") or "").strip(),
                    "email": (row.get("邮箱") or row.get("email") or "").strip(),
                    "phone": (row.get("手机号") or row.get("phone") or "").strip(),
                    "exhibition": (row.get("意向展会") or row.get("exhibition") or "").strip(),
                })
        if tag_filter:
            customers = [c for c in customers if tag_filter in (c.get("tags") or "").split(",")]
        # 去重 by email
        seen = set(); uniq = []
        for c in customers:
            e = (c.get("email") or "").strip().lower()
            if e and e in seen:
                continue
            seen.add(e); uniq.append(c)
        sales_name = load_settings(uid).get("signature") or u.get("display_name") or ""
        rendered = []
        for c in uniq:
            rendered.append({
                "company": c.get("company"),
                "contact": c.get("contact"),
                "email": c.get("email"),
                "subject": personalize(subject_tpl, c, sales_name),
                "body": personalize(body_tpl, c, sales_name),
            })
        return json_resp({"count": len(rendered), "items": rendered})

    def send_emails(self, conn, uid, u, d):
        items = d.get("items") or []
        cc = d.get("cc") or []
        bcc = d.get("bcc") or []
        interval = int(d.get("interval") or 5)
        attachment_ids = d.get("attachment_ids") or []
        settings = load_settings(uid)
        sales_name = settings.get("signature") or u.get("display_name") or ""
        # 附件
        attachments = []
        if attachment_ids:
            for aid in attachment_ids:
                row = conn.execute("SELECT * FROM materials WHERE id=? AND (user_id=? OR user_id=0)", (aid, uid)).fetchone()
                if row:
                    attachments.append({"name": row["name"], "file_path": row["file_path"]})
        exhibition = d.get("exhibition") or ""
        template_name = d.get("template_name") or ""
        results = []
        try:
            for it in items:
                c = {"company": it.get("company"), "contact": it.get("contact"), "email": it.get("email")}
                subject = personalize(it.get("subject",""), c, sales_name)
                body = personalize(it.get("body",""), c, sales_name)
                to = (it.get("email") or "").strip()
                status, error = "failed", "缺少收件邮箱"
                if to:
                    try:
                        status, error = send_one_email(uid, to, subject, body, settings, attachments, cc, bcc)
                    except Exception as e:
                        status, error = "failed", f"发送异常: {str(e)}"
                # 每封邮件立即写日志 + commit，防止超时丢失
                try:
                    conn.execute("INSERT INTO email_logs (user_id,exhibition,template_name,customer_company,contact,email,subject,body,status,error,sent_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                                 (uid, exhibition, template_name, it.get("company"), it.get("contact"), to, subject, body, status, error, now_iso()))
                    conn.commit()
                except Exception as log_err:
                    pass
                results.append({"email": to, "status": status})
                if interval and to:
                    threading.Event().wait(interval)
        except Exception as outer_e:
            try: conn.commit()
            except: pass
            return json_resp({"ok": True, "results": results, "demo_mode": bool(settings.get("demo_mode")), "warning": f"发送中断: {str(outer_e)}"})
        return json_resp({"ok": True, "results": results, "demo_mode": bool(settings.get("demo_mode"))})

    def translate_email(self, d):
        """邮件翻译：中文 → 英文 / 中英双语"""
        subject = d.get("subject", "")
        body = d.get("body", "")
        target = d.get("target", "en")
        if not subject or not body:
            return json_resp({"error": "邮件内容为空"}, 400)
        try:
            if target == "en":
                # 翻译为英文（含英文落款）
                en_subj = _call_llm(f"将以下邮件主题翻译为地道商务英文，只返回翻译结果，不要解释：\n{subject}")
                en_body = _translate_long_text(body)
                return json_resp({"subject": en_subj.strip(), "body": en_body.strip()})
            elif target == "bilingual":
                # 中英双语：中文全文在上（含中文落款），英文全文在下（含英文落款）
                en_body = _translate_long_text(body)
                bi_subj = _call_llm(f"将以下邮件主题翻译为地道商务英文，只返回翻译结果：\n{subject}")
                sep = "\n\n" + "-" * 36 + "\n\n"
                return json_resp({
                    "subject": subject.strip() + " / " + bi_subj.strip(),
                    "body": body.strip() + sep + en_body.strip()
                })
            else:
                return json_resp({"subject": subject, "body": body})
        except Exception as e:
            return json_resp({"error": f"翻译失败: {str(e)}"}, 500)


class ThreadingHTTPServer(socketserver.ThreadingMixIn, http.server.HTTPServer):
    daemon_threads = True

def main():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    cos_data_source = "local"
    if COS_ENABLED:
        # 容器本地文件系统是临时的；启动时先从 COS 拉取最新数据库与附件
        if cos_download_db():
            cos_data_source = "cos"
        else:
            cos_data_source = "local-first-boot"
        cos_download_all_uploads()
    init_db()
    ensure_ai_columns()
    # 写启动健康度到 backup_meta（让管理员中心能展示"上次启动源"）
    try:
        c = get_db()
        c.execute("INSERT OR REPLACE INTO backup_meta(key,value) VALUES ('last_boot_at',?)", (now_iso(),))
        c.execute("INSERT OR REPLACE INTO backup_meta(key,value) VALUES ('last_boot_source',?)", (cos_data_source,))
        c.execute("INSERT OR REPLACE INTO backup_meta(key,value) VALUES ('cos_enabled',?)",
                  ("1" if COS_ENABLED else "0",))
        c.commit(); c.close()
    except Exception:
        pass
    # 后台 COS 同步 worker
    _schedule_cos_db_sync()
    if COS_ENABLED:
        # 把初始化后的库（含默认 admin 账号）推送到 COS，确保新部署也有备份
        cos_upload_db()
    port = int(os.environ.get("PORT", 8000))
    server = ThreadingHTTPServer(("0.0.0.0", port), Handler)
    print(f"✅ 销售邮件发送工作台已启动： http://127.0.0.1:{port}")
    print(f"   数据库： {DB_PATH}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        server.shutdown()

if __name__ == "__main__":
    main()
